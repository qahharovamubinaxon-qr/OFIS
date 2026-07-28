"""Swapping the previous worker out of any firm's Word template.

Each firm words its договор differently — ГПХ or трудовой, one page or twelve —
but the worker's own data always carries the Госуслуги labels. These tests use
two firms whose wording has nothing else in common.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

from src.services import docx_worker as W  # noqa: E402

NEW = {
    "surname": "Назаров", "name": "Муродулло", "patronymic": "Хаиталиевич",
    "birth_date": "22.02.2004", "gender": "Мужской", "citizenship": "Узбекистан",
    "birth_place": "Узбекистан",
    "pass_series": "FB", "pass_number": "1234567",
    "pass_issue_date": "16.02.2023", "pass_issued_by": "МВД 99999",
    "pat_series": "77", "pat_number": "2600017664", "region": "Москва",
    "blank_series": "ПР", "blank_number": "4875056",
    "profession": "Штукатур", "contract_date": "28.07.2026",
    "work_address": "Москва",
}


def _doc(lines):
    d = docx.Document()
    for line in lines:
        d.add_paragraph(line)
    return d


def _text(doc) -> list[str]:
    return [" ".join(p.text.split()) for p in doc.paragraphs]


def _joined(doc) -> str:
    return "\n".join(_text(doc))


# ------------------------------------------------------- the уведомление


UVED = [
    "Сведения о работодателе",
    "Полное наименование",
    'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ "АКИТАС"',
    "ОГРН 1247700301133", "ИНН 7743447264", "КПП 774301001",
    "Сведения об иностранном гражданине",
    "Фамилия (рус.) Абдимуминов",
    "Имя (рус.) Хусниддин",
    "Отчество (рус.) Рашид Угли",
    "Дата рождения 31.12.1993",
    "Пол Мужской",
    "Гражданство", "", "Узбекистан",
    "Место рождения, населенный пункт Узбекистан",
    "Документ, удостоверяющий личность иностранного гражданина",
    "Вид документа Иностранный паспорт",
    "Серия", "FА", "",
    "Номер 4956294",
    "Дата выдачи 08.02.2022",
    "Кем выдан", "МВД 10237",
    "Сведения о разрешении на работу или патенте",
    "Документ Патент ИГ (ЛБГ)",
    "Серия 50",
    "Номер 2600089581",
    "Регион", "Московская область", "",
    "Серия бланка ПР",
    "Номер бланка 7430108",
    "Сведения о трудовой деятельности",
    "Профессия, специальность, должность, вид трудовой деятельности по договору",
    "Подсобный рабочий", "",
    "Нет подходящей профессии Нет",
    "Вид договора Трудовой",
    "Дата заключения договора 26.05.2026",
    "Адрес места работы", "Московская область",
    "Наименование территориального органа МВД России на региональном уровне",
    "ОПВМ ОМВД РОССИИ ПО Г.О. МЫТИЩИ МОСКОВСКОЙ ОБЛАСТИ",
]


def test_every_labelled_value_of_the_uvedomlenie_is_replaced() -> None:
    doc = _doc(UVED)
    report = W.swap_worker(doc, NEW)
    body = _joined(doc)

    assert not report.skipped
    for gone in ("Абдимуминов", "Хусниддин", "Рашид Угли", "31.12.1993",
                 "4956294", "08.02.2022", "МВД 10237", "2600089581",
                 "7430108", "26.05.2026", "Подсобный рабочий"):
        assert gone not in body, gone
    for value in NEW.values():
        assert value in body, value


def test_the_employers_own_block_is_never_touched() -> None:
    """The firm's name, ОГРН, ИНН and КПП belong to the template, not the worker."""
    doc = _doc(UVED)
    W.swap_worker(doc, NEW)
    body = _joined(doc)
    for kept in ('ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ "АКИТАС"',
                 "ОГРН 1247700301133", "ИНН 7743447264", "КПП 774301001",
                 "ОПВМ ОМВД РОССИИ ПО Г.О. МЫТИЩИ МОСКОВСКОЙ ОБЛАСТИ"):
        assert kept in body, kept


def test_polnoe_naimenovanie_is_not_read_as_the_gender_label() -> None:
    doc = _doc(["Полное наименование", "ООО «АКИТАС»"])
    report = W.swap_worker(doc, NEW)
    assert not report.filled
    assert _text(doc) == ["Полное наименование", "ООО «АКИТАС»"]


def test_a_value_the_form_prints_on_the_next_line_is_followed() -> None:
    """Госуслуги put «Гражданство», «Серия», «Кем выдан» above their values."""
    doc = _doc(UVED)
    W.swap_worker(doc, NEW)
    lines = _text(doc)
    assert lines[lines.index("Гражданство") + 2] == "Узбекистан"
    assert lines[lines.index("Кем выдан") + 1] == "МВД 99999"
    assert lines[lines.index("Регион") + 1] == "Москва"


def test_series_and_number_follow_the_section_they_sit_under() -> None:
    doc = _doc(UVED)
    W.swap_worker(doc, NEW)
    lines = _text(doc)
    assert lines[lines.index("Серия") + 1] == "FB"          # passport section
    assert "Номер 1234567" in lines
    assert "Серия 77" in lines                              # patent section
    assert "Номер 2600017664" in lines


def test_the_blank_series_label_wins_over_the_bare_one() -> None:
    doc = _doc(["Сведения о разрешении на работу или патенте",
                "Серия бланка ПР", "Номер бланка 7430108"])
    W.swap_worker(doc, NEW)
    assert _text(doc)[1:] == ["Серия бланка ПР", "Номер бланка 4875056"]


def test_a_label_the_caller_has_no_value_for_is_left_as_it_was() -> None:
    doc = _doc(["Дата рождения 31.12.1993", "Пол Мужской"])
    W.swap_worker(doc, {"gender": "Женский"})
    assert _text(doc) == ["Дата рождения 31.12.1993", "Пол Женский"]


def test_a_value_that_does_not_look_like_its_label_is_reported_not_mangled():
    """Better to leave a line alone than to write a date over a sentence."""
    doc = _doc(["Дата рождения не указана"])
    report = W.swap_worker(doc, NEW)
    assert _text(doc) == ["Дата рождения не указана"]
    assert "birth_date" in report.skipped


# ---------------------------------------------------------- the договор


DOGOVOR = [
    "ДОГОВОР ГПХ",
    "с физическим лицом на оказание услуг № 290355",
    "Московская область\t26 мая 2026 года",
    "Центр правовой помощи трудовым мигрантам и ООО «АКИТАС», ИНН 7743447264,",
    "Фамилия (рус.) Абдимумин Имя (рус.) Хусниддин Отчество (рус.) Рашид Угли",
    "Дата рождения 31.12.1993",
    "Пол Мужской Гражданство Узбекистан",
    "Иностранный паспорт Серия FA Номер 4956294 Дата выдачи 08.02.2022 "
    "Кем выдан МВД 10237",
    "Профессия, специальность, должность: Подсобный рабочий",
    "с другой стороны, именуемый(-ая) в дальнейшем Исполнитель",
    "Согласно Письму Министерства финансов от 24 февраля 2016 г. № 03-04-06/10104",
    "9. РЕКВИЗИТЫ СТОРОН.",
    "Юридический адрес: 141008, обл. Московская, г. Мытищи, ул. Мира, д. 37",
    "Генеральный директор ______________ А.В.Нуар.",
    "Ф.И.О Фамилия(рус.) Абдимуминов Имя (рус.) Хусниддин Отчество (рус.) Рашид Угли",
    "Адрес : Московская область, г.о. Выдное, рп Боброво",
]


def test_the_contracts_worker_block_is_swapped_line_by_line() -> None:
    doc = _doc(DOGOVOR)
    W.swap_worker(doc, NEW)
    lines = _text(doc)
    assert lines[4] == ("Фамилия (рус.) Назаров Имя (рус.) Муродулло "
                        "Отчество (рус.) Хаиталиевич")
    assert lines[7] == ("Иностранный паспорт Серия FB Номер 1234567 "
                        "Дата выдачи 16.02.2023 Кем выдан МВД 99999")
    assert lines[8] == "Профессия, специальность, должность: Штукатур"


def test_a_paspoprt_series_inline_is_not_taken_for_the_patents() -> None:
    """«Иностранный паспорт Серия … Номер …» says which document it is."""
    doc = _doc(["Иностранный паспорт Серия FA Номер 4956294"])
    W.swap_worker(doc, NEW)
    assert _text(doc) == ["Иностранный паспорт Серия FB Номер 1234567"]


def test_the_firms_legal_text_and_requisites_survive_untouched() -> None:
    doc = _doc(DOGOVOR)
    W.swap_worker(doc, NEW)
    lines = _text(doc)
    assert lines[3].endswith("ИНН 7743447264,")
    assert lines[11] == "9. РЕКВИЗИТЫ СТОРОН."
    assert lines[12].startswith("Юридический адрес: 141008")
    assert lines[13] == "Генеральный директор ______________ А.В.Нуар."


def test_the_previous_workers_home_address_does_not_travel_on() -> None:
    doc = _doc(DOGOVOR)
    W.swap_worker(doc, {**NEW, "address": ""})
    assert _text(doc)[15] == "Адрес :"


def test_a_home_address_that_is_known_is_written_in() -> None:
    doc = _doc(DOGOVOR)
    W.swap_worker(doc, {**NEW, "address": "г. Москва, ул. Мира, д. 1"})
    assert _text(doc)[15] == "Адрес : г. Москва, ул. Мира, д. 1"


# --------------------------------------------------------------- dating


def test_the_contract_header_is_re_dated() -> None:
    doc = _doc(DOGOVOR)
    assert W.swap_header_date(doc, "28 июля 2026") is not None
    assert _text(doc)[2] == "Московская область 28 июля 2026 года"


def test_the_firms_own_wording_after_the_date_is_kept() -> None:
    for tail, expected in (("года", "28 июля 2026 года"),
                           ("г.", "28 июля 2026 г."),
                           ("", "28 июля 2026")):
        doc = _doc([f"г. Москва 26 мая 2026 {tail}".strip()])
        W.swap_header_date(doc, "28 июля 2026")
        assert _text(doc)[0] == f"г. Москва {expected}"


def test_a_ministry_letter_deep_in_the_text_is_not_re_dated() -> None:
    """«от 24 февраля 2016 г.» is the law the firm cites, not the contract."""
    doc = _doc(DOGOVOR)
    W.swap_header_date(doc, "28 июля 2026")
    assert "24 февраля 2016 г." in _text(doc)[10]


def test_dating_stops_at_the_worker_block() -> None:
    doc = _doc(["ДОГОВОР", "Дата рождения 31.12.1993",
                "подписан 26 мая 2026 года"])
    assert W.swap_header_date(doc, "28 июля 2026") is None
    assert _text(doc)[2] == "подписан 26 мая 2026 года"


# ------------------------------------------------------------ mechanics


def test_a_value_keeps_the_font_the_firm_gave_it() -> None:
    """Rewriting the whole paragraph would put the label's font on the value."""
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("Дата рождения ")
    bold = p.add_run("31.12.1993")
    bold.bold = True
    bold.font.name = "Courier New"

    W.swap_worker(doc, NEW)
    runs = doc.paragraphs[0].runs
    assert [r.text for r in runs] == ["Дата рождения ", "22.02.2004"]
    assert runs[1].bold is True
    assert runs[1].font.name == "Courier New"
    assert runs[0].bold is not True


def test_a_value_split_across_many_runs_is_replaced_once() -> None:
    """Converted templates chop a value into a run per letter."""
    doc = docx.Document()
    p = doc.add_paragraph()
    for chunk in ("Дат", "а ", "рожде", "ния ", "31.", "12.", "1993"):
        p.add_run(chunk)
    W.swap_worker(doc, NEW)
    assert doc.paragraphs[0].text == "Дата рождения 22.02.2004"


def test_a_value_hidden_inside_a_hyperlink_is_still_found() -> None:
    """``paragraph.runs`` skips those, which is how a name survives a swap."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("Фамилия (рус.) ")
    p._p.append(parse_xml(
        f'<w:hyperlink {nsdecls("w")}><w:r><w:t>Абдимуминов</w:t></w:r></w:hyperlink>'))

    W.swap_worker(doc, NEW)
    assert W.text_of(W.runs_of(doc.paragraphs[0])) == "Фамилия (рус.) Назаров"


def test_values_inside_table_cells_are_reached() -> None:
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Дата рождения 31.12.1993"
    table.cell(0, 1).text = "Пол Мужской"
    W.swap_worker(doc, {**NEW, "gender": "Женский"})
    assert table.cell(0, 0).text == "Дата рождения 22.02.2004"
    assert table.cell(0, 1).text == "Пол Женский"


def test_a_run_holding_a_tab_beside_a_value_keeps_it() -> None:
    doc = _doc(["Московская область\t26 мая 2026 года"])
    W.swap_header_date(doc, "28 июля 2026")
    assert doc.paragraphs[0].text == "Московская область\t28 июля 2026 года"


# ----------------------------------------------------------- end to end


def test_the_service_fills_a_word_pair_for_a_firm_it_has_never_seen(tmp_path):
    from src.domain.documents import Passport, Patent
    from src.domain.enums import Gender
    from src.domain.trud_firm import TrudFirm
    from src.services.trud_service import TrudService

    trud, uved = tmp_path / "trudovoy.docx", tmp_path / "uvedomlenie.docx"
    _doc(DOGOVOR).save(str(trud))
    _doc(UVED).save(str(uved))
    firm = TrudFirm(name="АКИТАС", internal_code="akitas",
                    trud_template_path=trud, uved_template_path=uved)

    result = TrudService().generate(
        Passport(surname="НАЗАРОВ", name="МУРОДУЛЛО", patronymic="ХАИТАЛИЕВИЧ",
                 number="1234567", series="FB", birth_date=date(2004, 2, 22),
                 issue_date=date(2023, 2, 16), issued_by="МВД 99999",
                 nationality="УЗБЕКИСТАН", gender=Gender.MALE,
                 birth_place="УЗБЕКИСТАН"),
        Patent(series="77", number="2600017664", issue_date=date(2026, 4, 14),
               issued_by="ГУ МВД России по г. Москве", profession="Штукатур",
               blank_series="ПР", blank_number="4875056"),
        firm, form_date=date(2026, 7, 28), profession="Штукатур",
        output_dir=tmp_path / "out")

    for path in (result.trud_path, result.uved_path):
        assert path.suffix == ".docx", "no Word on this machine is fine"
        body = _joined(docx.Document(str(path)))
        assert "Назаров" in body and "Абдимуминов" not in body
        assert "1234567" in body and "4956294" not in body
        assert "Штукатур" in body and "Подсобный рабочий" not in body

    contract = _joined(docx.Document(str(result.trud_path)))
    assert "28 июля 2026 года" in contract
    assert "ИНН 7743447264" in contract           # the firm's own data survives
    assert "рп Боброво" not in contract           # the old worker's address does not


def test_the_uvedomlenie_keeps_the_firms_mvd_office(tmp_path) -> None:
    """It is chosen per firm when the firm is added, not per worker."""
    uved = tmp_path / "uvedomlenie.docx"
    _doc(UVED).save(str(uved))
    from src.services.docx_editor import TrudDocxEditor

    out = TrudDocxEditor().fill_uvedomlenie(uved, tmp_path / "o.docx", NEW)
    assert "ОПВМ ОМВД РОССИИ ПО Г.О. МЫТИЩИ МОСКОВСКОЙ ОБЛАСТИ" in _joined(
        docx.Document(str(out)))


def test_a_docx_pair_needs_no_study(tmp_path) -> None:
    """Only a PDF template has to be measured; Word is filled by text."""
    from src.common.errors import ValidationError
    from src.domain.trud_firm import TrudFirm
    from src.services.trud_service import TrudFirmService

    trud, uved = tmp_path / "trudovoy.docx", tmp_path / "uvedomlenie.docx"
    _doc(DOGOVOR).save(str(trud))
    _doc(UVED).save(str(uved))
    firm = TrudFirm(name="АКИТАС", internal_code="akitas",
                    trud_template_path=trud, uved_template_path=uved)
    service = TrudFirmService(repo=None)          # the refusal is before any I/O
    for study in (service.study_trud, service.study_uved):
        with pytest.raises(ValidationError):
            study(firm, object())
