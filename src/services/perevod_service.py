"""ПЕРЕВОД — notarial translation of personal documents into Russian.

The office's notary offers translation services: a client's passport, driving
licence, birth/marriage certificate, diploma or аттестат is translated into
Russian, printed on the office's own three pre-printed sheets and certified by
the notary.

Flow: drop the document photos (front/back) → the AI recognises WHICH document
it is, reads every field and returns a structured translation → the program
lays the package out on the three uploaded blanks:

* **sheet 1** — the original itself, turned black-and-white and centred;
* **sheet 2** — the translation;
* **sheet 3** — nothing at all. It is the notary's own certification sheet:
  he writes the registry number, the date, the names and the seal on it by
  hand, so the program must not print a word there.

All three come out as ONE three-page PDF (plus the translation as Word, for
correcting a name before printing).

``templates/perevod/forms.v1.json`` holds the canonical field order per
document type (CIS passports, driving licences, certificates …) so output stays
consistent no matter how the AI phrases things.
"""

from __future__ import annotations

import json
import shutil
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from src.ai.russian import RUSSIAN_RULES
from src.ai.text_client import ask
from src.common.errors import OfisError, ValidationError
from src.common.logging import get_logger
from src.config import paths
from src.domain.passport_rules import in_russian_letters
from src.pdf import perevod_pasport
from src.pdf.engine import _font_file
from src.pdf.perevod_spec import (
    A4_LONG,
    A4_SHORT,
    BLANK_STEMS,
    LABEL_SHARE,
    LEADING,
    MM,
    PASSPORT_PAGE_MM,
    PASSPORT_SPREAD_MM,
    REAL_MM,
    SCAN_BOX,
    SCAN_GAP_MM,
    SPREAD_ASPECT,
    TEXT_BOX,
    TEXT_OPACITY,
    TEXT_SIZE,
    TEXT_SIZE_MIN,
)

log = get_logger(__name__)

DOC_TYPES = [
    ("auto", "Авто (программа ўзи аниқлайди)"),
    ("passport", "Паспорт"),
    ("id_card", "АЙДИ карта (шахсга далолатнома)"),
    ("driver_license", "Ҳайдовчилик гувоҳномаси"),
    ("birth_certificate", "Туғилганлик ҳақида гувоҳнома"),
    ("marriage_certificate", "Никоҳ гувоҳномаси"),
    ("diploma", "Диплом"),
    ("attestat", "Аттестат"),
    ("migration_card", "Миграционная карта"),
    ("other", "Бошқа ҳужжат"),
]

#: What the office may hand over as a blank — its own sheet as PDF, or a scan
#: of it as a picture.
BLANK_SUFFIXES = (".pdf", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff")

_PROMPT = RUSSIAN_RULES + """Ты — присяжный переводчик, готовишь НОТАРИАЛЬНЫЙ перевод документа \
на русский язык.

На фотографиях один документ (может быть несколько сторон одного документа).
{hint}

ЗАДАЧА:
1. Определи вид документа и язык оригинала.
2. Прочитай ВСЕ поля документа, включая печати, штампы и надписи.
3. Переведи на русский язык по стандартам нотариального перевода:
   - имена собственные — транслитерация по правилам русского языка
     (например: Shahboz Isakov → Шахбоз Исаков);
   - даты — в формате ДД.ММ.ГГГГ;
   - названия органов и учреждений — полный официальный перевод;
   - если поле нечитаемо — значение "неразборчиво";
   - если поля нет в документе — не включай его вовсе.

Верни СТРОГО JSON:
{{
 "doc_type": "<passport|id_card|driver_license|birth_certificate|\
marriage_certificate|diploma|attestat|migration_card|other>",
 "source_language": "<язык оригинала в родительном падеже, напр. узбекского>",
 "title": "<название документа по-русски заглавными, напр. ПАСПОРТ>",
 "issuing_country": "<государство, выдавшее документ>",
 "fields": [{{"label": "<поле по-русски>", "value": "<перевод значения>"}}],
 "stamps": ["<перевод текста печатей и штампов, если есть>"],
 "notes": ["<примечания переводчика, если нужны>"],
 "crops": [{{"x0": <число>, "y0": <число>, "x1": <число>, "y1": <число>}}]
}}

Если это ПАСПОРТ или ID-карта, поля должны называться РОВНО так: "Тип", \
"Код государства", "Номер паспорта", "Фамилия", "Имя", "Отчество", \
"Гражданство", "Дата рождения", "Место рождения", "Пол", "Дата выдачи", \
"Действителен до", "Орган, выдавший документ" (у ID-карты ещё "ПИНФЛ"). \
"Пол" — одна буква: "Ж" или "М". "Место рождения" — по-русски \
(SURKHANDARYA REGION → СУРХАНДАРЬИНСКАЯ ОБЛАСТЬ). "Орган, выдавший \
документ" КОПИРУЕТСЯ как напечатано, только русскими буквами, вместе с \
номером отделения: "MIA 22204" → "МВД 22204"; не расшифровывай его.

Поле "crops" — по одному элементу на КАЖДУЮ присланную фотографию, в том же \
порядке. Это границы самого документа на фото (без стола, кровати, пальцев и \
прочего фона) в долях от размера изображения: 0.0 — левый/верхний край, 1.0 — \
правый/нижний. Если документ занимает всё фото, верни 0,0,1,1.
Только JSON, без пояснений и без markdown."""

@dataclass(frozen=True)
class PerevodResult:
    pdf_path: Path
    docx_path: Path
    doc_type: str
    title: str


# ---------------------------------------------------------------- blanks


def blanks_dir() -> Path:
    """Where the office's own three sheets live.

    In AppData, never in the program folder — rebuilding the EXE must not throw
    the office's blanks away.
    """
    folder = paths.user_templates_dir() / "perevod"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def blank_path(index: int) -> Path | None:
    """The uploaded blank for sheet ``index`` (1..3), or None if there is none."""
    if not 1 <= index <= len(BLANK_STEMS):
        return None
    stem = BLANK_STEMS[index - 1]
    for suffix in BLANK_SUFFIXES:
        candidate = blanks_dir() / f"{stem}{suffix}"
        if candidate.exists():
            return candidate
    return None


def blanks() -> list[Path | None]:
    return [blank_path(i) for i in range(1, len(BLANK_STEMS) + 1)]


def set_blank(index: int, source: Path) -> Path:
    """Register one sheet. Replaces whatever was there under any suffix."""
    if not 1 <= index <= len(BLANK_STEMS):
        raise ValidationError("Бланка рақами 1, 2 ёки 3 бўлиши керак",
                              context={"index": index})
    source = Path(source)
    suffix = source.suffix.lower()
    if suffix not in BLANK_SUFFIXES:
        raise ValidationError(
            "Бланка PDF ёки расм бўлиши керак (pdf, png, jpg)",
            context={"path": str(source)})
    if not source.exists():
        raise ValidationError("Бланка файли топилмади", context={"path": str(source)})
    clear_blank(index)
    dest = blanks_dir() / f"{BLANK_STEMS[index - 1]}{suffix}"
    shutil.copyfile(source, dest)
    log.info("ПЕРЕВОД: %d-бланка юкланди — %s", index, dest.name)
    return dest


def clear_blank(index: int) -> None:
    existing = blank_path(index)
    if existing is not None:
        existing.unlink(missing_ok=True)


# ------------------------------------------------------------- the emblem


def emblem_path() -> Path | None:
    """The state emblem the drawn sheet puts at the head of the frame.

    Uploaded once by the office (the same picture stands at the top of every
    Uzbek passport), kept in AppData beside the blanks. Without it the frame
    is drawn all the same — only that window stays empty.
    """
    for suffix in BLANK_SUFFIXES:
        candidate = blanks_dir() / f"emblem{suffix}"
        if candidate.exists():
            return candidate
    return None


def set_emblem(source: Path) -> Path:
    source = Path(source)
    if source.suffix.lower() not in BLANK_SUFFIXES:
        raise ValidationError("Герб расм бўлиши керак (png, jpg)",
                              context={"path": str(source)})
    if not source.exists():
        raise ValidationError("Герб файли топилмади", context={"path": str(source)})
    clear_emblem()
    dest = blanks_dir() / f"emblem{source.suffix.lower()}"
    shutil.copyfile(source, dest)
    log.info("ПЕРЕВОД: герб юкланди — %s", dest.name)
    return dest


def clear_emblem() -> None:
    existing = emblem_path()
    if existing is not None:
        existing.unlink(missing_ok=True)


def _emblem_bytes() -> bytes | None:
    found = emblem_path()
    try:
        return found.read_bytes() if found is not None else None
    except OSError:
        return None


# ------------------------------------------------------------ the original


def photocopy(image: bytes, crop: dict | None = None) -> bytes:
    """The original as a clean black-and-white copy, ready to be pasted on.

    Colourless — grey ink on white paper, the way a photocopier gives it. NOT a
    negative: the office tried that and a negative turns the whole sheet black
    and eats a printer's toner, and a passport copy is read from the black.

    ``crop`` are the document's bounds within the photo in 0..1 fractions (the
    AI returns them alongside the translation — it can see where the document
    is far more reliably than edge detection can on a patterned surface). The
    photo is trimmed to those bounds, its lighting evened out and lifted to
    paper-white / ink-black.

    Returns the original bytes untouched if OpenCV is unavailable or the bytes
    are not a readable image — a package must still come out.
    """
    import numpy as np

    try:
        import cv2
    except ImportError:  # pragma: no cover - cv2 ships with the app
        return image

    arr = cv2.imdecode(np.frombuffer(image, np.uint8), cv2.IMREAD_COLOR)
    if arr is None:
        return image
    arr = _apply_crop(arr, crop)
    out = _scan_look(cv2, arr)
    return out or image


def _apply_crop(arr, crop: dict | None):
    """Trim to the AI-reported document bounds, with a small safety margin."""
    if not isinstance(crop, dict):
        return arr
    try:
        x0 = float(crop["x0"]); y0 = float(crop["y0"])
        x1 = float(crop["x1"]); y1 = float(crop["y1"])
    except (KeyError, TypeError, ValueError):
        return arr
    if not (0.0 <= x0 < x1 <= 1.0 and 0.0 <= y0 < y1 <= 1.0):
        return arr
    if (x1 - x0) * (y1 - y0) < 0.05:  # implausibly small — ignore
        return arr
    h, w = arr.shape[:2]
    pad = 0.012
    px0 = max(0, int((x0 - pad) * w)); py0 = max(0, int((y0 - pad) * h))
    px1 = min(w, int((x1 + pad) * w)); py1 = min(h, int((y1 + pad) * h))
    if px1 - px0 < 40 or py1 - py0 < 40:
        return arr
    return arr[py0:py1, px0:px1]


def _scan_look(cv2, arr) -> bytes | None:
    """Flatten phone lighting and lift the page to crisp black-on-white."""
    gray = cv2.cvtColor(arr, cv2.COLOR_BGR2GRAY)
    # divide out the illumination field, then stretch what is left
    blur = cv2.GaussianBlur(gray, (0, 0), sigmaX=max(gray.shape) / 30.0)
    flat = cv2.divide(gray, blur, scale=255)
    flat = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8)).apply(flat)
    # clip the paper to white and the ink to black without going fully binary,
    # so photos, stamps and guilloche stay readable
    lo, hi = 120.0, 225.0
    stretched = cv2.convertScaleAbs(flat, alpha=255.0 / (hi - lo),
                                    beta=-lo * 255.0 / (hi - lo))
    ok, buf = cv2.imencode(".png", stretched)
    return buf.tobytes() if ok else None


# -------------------------------------------------------------- the sheets


def _blank_page(doc, blank: Path | None):
    """Start a sheet on ``blank``, keeping the blank's own proportions.

    A PDF blank is inserted as-is, so its text, rules and letterhead stay
    exactly as the office printed them. A picture blank gets a page cut to the
    picture's own shape (long side A4) and is laid over the whole of it — no
    stretching, no white bands. With no blank at all the sheet is plain A4, so
    the section still works the day before the office uploads its sheets.
    """
    import fitz

    if blank is not None and blank.suffix.lower() == ".pdf":
        try:
            source = fitz.open(str(blank))
        except Exception as exc:                      # noqa: BLE001
            log.warning("ПЕРЕВОД: бланка очилмади (%s): %s", blank.name, exc)
        else:
            with source:
                if source.page_count:
                    doc.insert_pdf(source, from_page=0, to_page=0)
                    return doc[-1]
            log.warning("ПЕРЕВОД: бланка бўш — %s", blank.name)

    width, height = A4_SHORT, A4_LONG
    stream = None
    if blank is not None:
        try:
            picture = fitz.Pixmap(str(blank))
            if picture.width and picture.height:
                if picture.width > picture.height:
                    width = A4_LONG
                    height = A4_LONG * picture.height / picture.width
                else:
                    height = A4_LONG
                    width = A4_LONG * picture.width / picture.height
                stream = blank
        except Exception as exc:                      # noqa: BLE001
            log.warning("ПЕРЕВОД: бланка расми ўқилмади (%s): %s", blank.name, exc)

    page = doc.new_page(width=width, height=height)
    if stream is not None:
        try:
            page.insert_image(page.rect, filename=str(stream), keep_proportion=False)
        except (RuntimeError, ValueError) as exc:
            log.warning("ПЕРЕВОД: бланка жойлашмади (%s): %s", blank.name, exc)
    return page


def _aspect(shot: bytes) -> float:
    """The photo's width over its height, 1.0 if it cannot be read."""
    import fitz

    try:
        picture = fitz.Pixmap(shot)
    except Exception:                                 # noqa: BLE001
        return 1.0
    return (picture.width / picture.height) if picture.height else 1.0


def real_size_pt(doc_type: str, aspect: float) -> tuple[float, float]:
    """How big this document is in real life, in points, for THIS photograph.

    The size comes from the document, the ORIENTATION from the photograph: a
    plastic card photographed on its side is still 85.6 × 54 mm, just turned.
    A passport is the one that changes size rather than shape — photographed
    open, the spread is two pages wide.
    """
    long_mm, short_mm = REAL_MM.get(doc_type, REAL_MM["other"])
    if doc_type == "passport":
        long_mm, short_mm = (PASSPORT_SPREAD_MM if aspect >= SPREAD_ASPECT
                             else PASSPORT_PAGE_MM)
    wide = long_mm >= short_mm
    if (aspect >= 1.0) != wide:
        long_mm, short_mm = short_mm, long_mm
    return long_mm * MM, short_mm * MM


def _place_originals(page, shots: list[bytes], doc_type: str) -> None:
    """The copies of the original, at LIFE SIZE, centred on the sheet.

    The office asked for this in as many words: a passport is not blown up to
    fill an A4, it is printed the size a passport is and put in the middle of
    the sheet — «ҳисоблаб ўртасига қўйилади». So each photograph is scaled to
    the document's own real size (:func:`real_size_pt`) rather than to the
    window, and the whole group is centred in the window both ways.

    Front and back of the same card both go on this sheet, one under the other.
    Only if the stack will not fit the window is it scaled down — evenly, so the
    two halves of one document stay the same size as each other.
    """
    import fitz

    if not shots:
        return
    rect = page.rect
    left = SCAN_BOX[0] * rect.width
    right = SCAN_BOX[2] * rect.width
    top = SCAN_BOX[1] * rect.height
    bottom = SCAN_BOX[3] * rect.height
    room_w, room_h = right - left, bottom - top
    gap = SCAN_GAP_MM * MM

    sizes = [real_size_pt(doc_type, _aspect(shot)) for shot in shots]
    stack_h = sum(h for _w, h in sizes) + gap * (len(sizes) - 1)
    stack_w = max(w for w, _h in sizes)
    # life size unless the sheet is too small for it
    fit = min(1.0, room_w / stack_w, room_h / stack_h)
    if fit < 1.0:
        log.info("ПЕРЕВОД: ҳақиқий ўлчам варақга сиғмади — %.0f%% қилинди",
                 fit * 100)

    y = top + (room_h - stack_h * fit) / 2.0
    centre_x = (left + right) / 2.0
    for shot, (width, height) in zip(shots, sizes, strict=True):
        width, height = width * fit, height * fit
        box = fitz.Rect(centre_x - width / 2.0, y,
                        centre_x + width / 2.0, y + height)
        try:
            page.insert_image(box, stream=shot, keep_proportion=True)
        except Exception as exc:                      # noqa: BLE001
            log.warning("ПЕРЕВОД: ҳужжат расми жойлашмади: %s", exc)
        y += height + gap * fit


def _forms() -> dict:
    path = paths.templates_dir() / "perevod" / "forms.v1.json"
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def _russian(field: dict) -> dict:
    """One read field, spelled the way a Russian document spells it.

    A Tajik document is translated by the reader but its own alphabet often
    survives in a name or a place — Хоҷа, Ғафуров, Қӯрғонтеппа. A notarial
    translation into Russian may not carry those letters, so the office's
    rule applies here too: Ҷ ҷ are ДЖ дж, and the rest drop to their nearest
    Russian letter.
    """
    return {"label": in_russian_letters(str(field.get("label", ""))),
            "value": in_russian_letters(str(field.get("value", "") or ""))}


def _order_fields(doc_type: str, fields: list[dict]) -> list[dict]:
    """Sort the AI's fields into the canonical order for this document type;
    anything unknown keeps its original position at the end."""
    template = _forms().get(doc_type) or {}
    wanted = [w.lower() for w in template.get("fields", [])]
    if not wanted:
        return fields

    def rank(f: dict) -> int:
        label = str(f.get("label", "")).strip().lower()
        for i, w in enumerate(wanted):
            if label == w or label.startswith(w) or w.startswith(label):
                return i
        return len(wanted) + 1

    return sorted(fields, key=rank)


class PerevodService:
    def __init__(self, key_getter, cert_getter=None) -> None:
        self._key_getter = key_getter
        # Kept for API compatibility only. The notary's names, the date and the
        # registry number are NOT printed any more: sheet 3 is his own blank and
        # he completes it in person.
        self._cert_getter = cert_getter

    # ---------------------------------------------------------- templates
    @staticmethod
    def blanks() -> list[Path | None]:
        return blanks()

    @staticmethod
    def set_blank(index: int, source: Path) -> Path:
        return set_blank(index, source)

    @staticmethod
    def clear_blank(index: int) -> None:
        clear_blank(index)

    @staticmethod
    def emblem() -> Path | None:
        return emblem_path()

    @staticmethod
    def set_emblem(source: Path) -> Path:
        return set_emblem(source)

    @staticmethod
    def clear_emblem() -> None:
        clear_emblem()

    # ----------------------------------------------------------- printing
    def translate(
        self,
        images: list[bytes],
        *,
        doc_type: str = "auto",
        form_date: date | None = None,  # kept for API compatibility; unused
        output_dir: Path | None = None,
    ) -> PerevodResult:
        if not images:
            raise OfisError("Hujjat rasmini yuklang.")
        hint = ""
        if doc_type != "auto":
            label = dict(DOC_TYPES).get(doc_type, doc_type)
            hint = f"Оператор указал вид документа: {label} ({doc_type}).\n"

        from src.ocr.preprocess import prepare_image

        prepared = [prepare_image(i) for i in images[:10]]
        raw = ask(self._key_getter(), _PROMPT.format(hint=hint),
                  prepared, json_out=True)
        try:
            data = json.loads(raw)
        except json.JSONDecodeError as exc:
            raise OfisError("AI javobini o'qib bo'lmadi — qayta urinib ko'ring.") from exc

        kind = str(data.get("doc_type") or (doc_type if doc_type != "auto" else "other"))
        title = str(data.get("title") or _forms().get(kind, {}).get("title") or "ДОКУМЕНТ")
        lang = str(data.get("source_language") or "иностранного")
        country = str(data.get("issuing_country") or "")
        fields = _order_fields(kind, [_russian(f) for f in data.get("fields", [])
                                      if isinstance(f, dict) and f.get("label")])
        stamps = [in_russian_letters(str(s)) for s in data.get("stamps", [])
                  if str(s).strip()]
        notes = [in_russian_letters(str(n)) for n in data.get("notes", [])
                 if str(n).strip()]
        crops = data.get("crops") if isinstance(data.get("crops"), list) else []

        folder = output_dir if output_dir is not None else paths.output_dir() / "perevod"
        folder.mkdir(parents=True, exist_ok=True)
        stem = self._stem(fields, title)
        base = folder / stem
        i = 1
        while base.with_suffix(".pdf").exists():
            base = folder / f"{stem}_{i:03d}"
            i += 1

        originals = [
            photocopy(shot, crop)
            for shot, crop in zip(images[:10],
                                  list(crops) + [None] * len(images),
                                  strict=False)
        ]
        pdf = self._to_pdf(base.with_suffix(".pdf"), title=title, lang=lang,
                           country=country, fields=fields, stamps=stamps,
                           notes=notes, originals=originals, doc_type=kind)
        docx = self._to_docx(base.with_suffix(".docx"), title=title, lang=lang,
                             country=country, fields=fields, stamps=stamps,
                             notes=notes)
        log.info("Perevod: %s (%s) → %s", title, kind, pdf.name)
        return PerevodResult(pdf_path=pdf, docx_path=docx, doc_type=kind, title=title)

    # ------------------------------------------------------------------
    @staticmethod
    def _stem(fields: list[dict], title: str) -> str:
        surname = ""
        for f in fields:
            label = str(f.get("label", "")).lower()
            if "фамилия" in label:
                surname = str(f.get("value", "")).split()[0] if f.get("value") else ""
                break
        raw = f"{surname}_{title}".strip("_") or "PEREVOD"
        return "".join(c if c.isalnum() or c in " _-" else "_" for c in raw).strip()

    def _to_pdf(self, out: Path, *, title: str, lang: str, country: str,
                fields: list[dict], stamps: list[str], notes: list[str],
                originals: list[bytes], doc_type: str = "other") -> Path:
        """The package: three sheets, one PDF, in the office's own order."""
        import fitz

        serif = fitz.Font(fontfile=str(_font_file("OfisSerif")))
        bold = fitz.Font(fontfile=str(_font_file("OfisSerifBold")))
        sheets = blanks()

        doc = fitz.open()

        # sheet 1 — the original, black-and-white, centred on the blank
        _place_originals(_blank_page(doc, sheets[0]), originals, doc_type)

        # sheet 2 — the translation. A passport or an ID card is DRAWN, the
        # way the office types it: an UZBEK one on the office's own sheet
        # (that pattern is the Uzbek passport's and nobody else's), any
        # other republic's as its own international data page. Anything
        # else is set as the «поле: значение» list below.
        page = _blank_page(doc, sheets[1])
        if doc_type in ("passport", "id_card"):
            drawn = perevod_pasport.from_fields(
                fields, lang=lang, country=country, title=title,
                stamps=stamps, notes=notes, emblem=_emblem_bytes())
            if perevod_pasport.is_drawable(drawn):
                if perevod_pasport.is_uzbek(drawn):
                    perevod_pasport.draw(page, drawn)
                else:
                    perevod_pasport.draw_generic(page, drawn)
                _blank_page(doc, sheets[2])
                out.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(out), garbage=4, deflate=True)
                doc.close()
                return out
            log.warning("ПЕРЕВОД: %s ўқилмади — рўйхат кўринишида чиқарилди",
                        doc_type)
        rect = page.rect
        x0 = TEXT_BOX[0] * rect.width
        width = (TEXT_BOX[2] - TEXT_BOX[0]) * rect.width
        top = TEXT_BOX[1] * rect.height
        room = (TEXT_BOX[3] - TEXT_BOX[1]) * rect.height

        size = TEXT_SIZE * rect.height
        floor = TEXT_SIZE_MIN * rect.height
        ops, height = _compose(title=title, lang=lang, country=country,
                               fields=fields, stamps=stamps, notes=notes,
                               serif=serif, bold=bold, width=width, size=size)
        while height > room and size > floor:
            size -= max(0.25, size * 0.04)
            ops, height = _compose(title=title, lang=lang, country=country,
                                   fields=fields, stamps=stamps, notes=notes,
                                   serif=serif, bold=bold, width=width, size=size)
        if height > room:
            log.warning("ПЕРЕВОД: таржима бир варақдан узун — энг кичик ўлчамда "
                        "чиқарилди (%.0f > %.0f)", height, room)

        writer = fitz.TextWriter(rect)
        for dx, dy, text, font, font_size in ops:
            writer.append((x0 + dx, top + dy), text, font=font, fontsize=font_size)
        writer.write_text(page, opacity=TEXT_OPACITY)

        # sheet 3 — the notary's own certification blank. NOTHING is printed on
        # it: no names, no city, no date, no registry number. He fills it and
        # seals it by hand, and a program-printed name on it would be a forgery.
        _blank_page(doc, sheets[2])

        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out), garbage=4, deflate=True)
        doc.close()
        return out

    @staticmethod
    def _to_docx(out: Path, *, title: str, lang: str, country: str,
                 fields: list[dict], stamps: list[str], notes: list[str]) -> Path:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        d = docx.Document()
        style = d.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        head = d.add_paragraph()
        run = head.add_run(f"ПЕРЕВОД С {lang.upper()} ЯЗЫКА НА РУССКИЙ ЯЗЫК")
        run.bold = True
        run.font.size = Pt(12)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if country:
            p = d.add_paragraph(country.upper())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t = d.add_paragraph()
        tr = t.add_run(title.upper())
        tr.bold = True
        tr.font.size = Pt(13)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.add_paragraph()

        for f in fields:
            label = str(f.get("label", "")).strip()
            if not label:
                continue
            p = d.add_paragraph()
            p.add_run(f"{label}: ")
            p.add_run(str(f.get("value", "")).strip()).bold = True

        if stamps:
            d.add_paragraph()
            d.add_paragraph("Печати и штампы:").runs[0].bold = True
            for s in stamps:
                d.add_paragraph(s)
        for n in notes:
            d.add_paragraph(f"Примечание переводчика: {n}")
        d.save(str(out))
        return out


# ------------------------------------------------------------- typesetting


def _wrap(text: str, font, size: float, avail: float) -> list[str]:
    """Break ``text`` into lines no wider than ``avail``.

    A single word longer than the line (a 30-character institution name run
    together by the reader) is cut rather than allowed to run off the sheet.
    """
    rows: list[str] = []
    current = ""
    for word in text.split():
        candidate = (current + " " + word).strip()
        if current and font.text_length(candidate, fontsize=size) > avail:
            rows.append(current)
            current = word
        else:
            current = candidate
        while font.text_length(current, fontsize=size) > avail and len(current) > 1:
            cut = len(current) - 1
            while cut > 1 and font.text_length(current[:cut], fontsize=size) > avail:
                cut -= 1
            rows.append(current[:cut])
            current = current[cut:]
    if current:
        rows.append(current)
    return rows


def _compose(*, title: str, lang: str, country: str, fields: list[dict],
             stamps: list[str], notes: list[str], serif, bold,
             width: float, size: float) -> tuple[list[tuple], float]:
    """The translation as draw operations relative to the block's top-left.

    Returns the operations and the height they need, so the caller can step the
    type size down until the whole translation fits the blank on one sheet.
    """
    ops: list[tuple] = []
    y = 0.0
    lead = size * LEADING
    label_width = width * LABEL_SHARE

    def centre(text: str, font, font_size: float) -> None:
        nonlocal y
        line_width = font.text_length(text, fontsize=font_size)
        ops.append(((width - line_width) / 2, y + font_size, text, font, font_size))
        y += font_size * LEADING

    def left(text: str, font, font_size: float, dx: float = 0.0) -> None:
        nonlocal y
        ops.append((dx, y + font_size, text, font, font_size))
        y += font_size * LEADING

    centre(f"ПЕРЕВОД С {lang.upper()} ЯЗЫКА НА РУССКИЙ ЯЗЫК", bold, size * 1.09)
    y += size * 0.7
    if country:
        centre(country.upper(), serif, size)
    centre(title.upper(), bold, size * 1.18)
    y += size * 1.1

    for f in fields:
        label = str(f.get("label", "")).strip()
        value = str(f.get("value", "")).strip()
        if not label:
            continue
        rows = _wrap(value, bold, size, width - label_width) or [""]
        for index, row in enumerate(rows):
            if index == 0:
                # a label wider than its column is set smaller rather than
                # allowed to run under the value beside it
                label_size = size
                text = f"{label}:"
                while (label_size > 4.0
                       and serif.text_length(text, fontsize=label_size)
                       > label_width - size * 0.3):
                    label_size -= 0.3
                ops.append((0.0, y + size, text, serif, label_size))
            ops.append((label_width, y + size, row, bold, size))
            y += lead

    if stamps:
        y += size
        centre("Печати и штампы:", bold, size)
        for stamp in stamps:
            for row in _wrap(stamp, serif, size, width):
                left(row, serif, size)

    if notes:
        y += size
        small = size * 0.9
        for note in notes:
            for row in _wrap(f"Примечание переводчика: {note}", serif, small, width):
                left(row, serif, small)

    return ops, y
