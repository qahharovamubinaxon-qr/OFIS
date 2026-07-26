"""ДОВЕРЕННОСТЬ: draft notarial documents for the office's notary partner.

The operator drops the parties' passport photos, picks a date and a document
type (or just describes it — the AI infers the type), and writes who → whom →
for what. Gemini composes the full Russian text per Moscow notarial practice,
addressed for certification by the office's notary; the draft is saved as BOTH
.docx and .pdf. The notary reviews, signs and stamps the print-out.
"""

from __future__ import annotations

import json
import urllib.request
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from src.common.errors import OfisError
from src.common.logging import get_logger
from src.config import paths
from src.domain.documents import Passport
from src.pdf.formatters import _date_dmy

log = get_logger(__name__)

NOTARY_FIO = "Друганова Маргарита Владимировна"
NOTARY_SHORT = "Друганова М.В."
NOTARY_CITY = "город Москва"

# settings keys (all editable in Sozlamalar)
KEY_SERIES_PREFIX = "dover.series_prefix"   # «77 АВ»
KEY_SERIES_NEXT = "dover.series_next"       # 2463964 → +1 per document
KEY_REESTR_NEXT = "dover.reestr_next"       # 12855 → +1 per document
KEY_TARIF = "dover.tarif"                   # «1500»
DEFAULT_SERIES_PREFIX = "77 АВ"
DEFAULT_SERIES_NEXT = 2463964
DEFAULT_REESTR_NEXT = 12855
DEFAULT_TARIF = "1500"

DOVER_TYPES = [
    "Авто (тавсифдан аниқлайди)",
    "Генеральная доверенность",
    "Доверенность на автомобиль",
    "Доверенность на ведение дел и представительство",
    "Доверенность на получение документов",
    "Доверенность на получение денежных средств",
    "Доверенность на распоряжение недвижимостью",
    "Доверенность на регистрационные действия",
    "Согласие на выезд ребёнка за границу",
    "Согласие на сопровождение ребёнка",
    "Согласие супруга на сделку",
    "Заявление (свободная форма)",
]

_SYSTEM = (
    "Ты — помощник московского нотариуса. Составь ПОЛНЫЙ текст нотариального "
    "документа (доверенность / согласие / заявление) по стандартам нотариальной "
    "практики города Москвы и законодательства РФ (ГК РФ, Основы законодательства "
    "о нотариате). Пиши строго официальным русским языком.\n"
    "СТРОГАЯ СТРУКТУРА (каждый пункт — с новой строки):\n"
    "1) первая строка — ТОЛЬКО вид документа ЗАГЛАВНЫМИ (СОГЛАСИЕ / "
    "ДОВЕРЕННОСТЬ / ЗАЯВЛЕНИЕ);\n"
    "2) вторая строка — «Город Москва.»;\n"
    "3) третья строка — дата составления ПРОПИСЬЮ, например «Двадцать шестое "
    "июля две тысячи двадцать шестого года.»;\n"
    "4) далее абзацы: полные данные доверителя (ФИО, дата рождения, пол, "
    "гражданство, паспорт: серия, номер, кем и когда выдан), полные данные "
    "представителя, подробные полномочия по смыслу задания, срок действия "
    "(если уместен — один год, если не указан иной), право/запрет передоверия, "
    "отметка о разъяснении статей закона;\n"
    "5) строка «Подпись:» с длинной линией из подчёркиваний;\n"
    "6) удостоверительная надпись: снова «Город Москва.», дата прописью на "
    "отдельной строке, затем «Настоящее согласие (доверенность/заявление) "
    "удостоверено мной, {notary}, нотариусом города Москвы.» и «…подписано "
    "гражданином <ФИО> в моем присутствии. Личность его установлена, "
    "дееспособность проверена.»\n"
    "НЕ пиши строки про реестровый номер, тариф и подпись нотариуса — их "
    "добавляет программа. Верни ТОЛЬКО текст документа, без пояснений и без "
    "markdown."
).format(notary=NOTARY_FIO)


@dataclass(frozen=True)
class DoverResult:
    docx_path: Path
    pdf_path: Path
    series: str = ""
    reestr: int = 0


def _passport_block(label: str, p: Passport | None) -> str:
    if p is None:
        return f"{label}: не указан"
    parts = [f"{label}: {p.surname} {p.name} {p.patronymic or ''}".strip()]
    if p.birth_date:
        parts.append(f"дата рождения {_date_dmy(p.birth_date)}")
    if p.nationality:
        parts.append(f"гражданство {p.nationality}")
    num = f"{p.series or ''}{p.number}".strip()
    parts.append(f"паспорт {num}")
    if p.issue_date:
        parts.append(f"выдан {_date_dmy(p.issue_date)}")
    if p.issued_by:
        parts.append(f"кем выдан: {p.issued_by}")
    return ", ".join(parts)


class DoverService:
    def __init__(self, key_getter, settings=None) -> None:
        self._key_getter = key_getter
        self._settings = settings  # SettingsService | None (counters persist when set)

    # -- counters ------------------------------------------------------
    def _counter(self, key: str, default: int) -> int:
        if self._settings is None:
            return default
        try:
            return int(self._settings.get(key, default) or default)
        except (TypeError, ValueError):
            return default

    def _str_setting(self, key: str, default: str) -> str:
        if self._settings is None:
            return default
        return str(self._settings.get(key, default) or default).strip() or default

    def _advance_counters(self) -> None:
        if self._settings is None:
            return
        self._settings.set(KEY_SERIES_NEXT,
                           self._counter(KEY_SERIES_NEXT, DEFAULT_SERIES_NEXT) + 1)
        self._settings.set(KEY_REESTR_NEXT,
                           self._counter(KEY_REESTR_NEXT, DEFAULT_REESTR_NEXT) + 1)

    def generate_from_images(
        self,
        images: list[bytes],
        *,
        doc_type: str,
        description: str,
        form_date: date,
        output_dir: Path | None = None,
    ) -> DoverResult:
        """Compose straight from the dropped document photos (10-15 allowed:
        passports, СТС front/back, …) — Gemini reads them itself."""
        text = self._compose_images(images, doc_type=doc_type,
                                    description=description, form_date=form_date)
        return self._save(text, "DOVER", output_dir)

    def _compose_images(self, images, *, doc_type, description, form_date) -> str:
        import base64

        key = (self._key_getter() or "").strip()
        if not key:
            raise OfisError("AI kaliti yo'q — Sozlamalarga Gemini kalitini kiriting.")
        from src.ocr.preprocess import prepare_image

        chosen = "" if doc_type.startswith("Авто") else f"Вид документа: {doc_type}. "
        user = (
            f"{chosen}Дата составления: {_date_dmy(form_date)}. "
            "Данные сторон и объекта (авто и т.п.) возьми из приложенных фото "
            "документов (паспорта, СТС и др.); нечитаемое оставь как «________». "
            f"Задание от оператора (кто, кому, для чего): {description or 'не указано'}"
        )
        parts = [{"text": _SYSTEM + "\n\n" + user}]
        for img in images[:15]:
            parts.append({"inline_data": {
                "mime_type": "image/jpeg",
                "data": base64.b64encode(prepare_image(img)).decode(),
            }})
        return self._call(key, parts)

    def _save(self, text: str, stem: str, output_dir: Path | None) -> DoverResult:
        from src.pdf.dover_renderer import finalize_notarial_text, render_dover_pdf

        reestr = self._counter(KEY_REESTR_NEXT, DEFAULT_REESTR_NEXT)
        tarif = self._str_setting(KEY_TARIF, DEFAULT_TARIF)
        series = (f"{self._str_setting(KEY_SERIES_PREFIX, DEFAULT_SERIES_PREFIX)} "
                  f"{self._counter(KEY_SERIES_NEXT, DEFAULT_SERIES_NEXT)}")
        final = finalize_notarial_text(text, reestr=reestr, tarif=tarif,
                                       notary_short=NOTARY_SHORT)

        title = next((ln.strip() for ln in final.splitlines() if ln.strip()),
                     "ДОВЕРЕННОСТЬ")
        kind = "".join(c for c in title.split()[0] if c.isalpha()) or "ДОВЕРЕННОСТЬ"
        folder = output_dir if output_dir is not None else paths.output_dir() / "dover"
        folder.mkdir(parents=True, exist_ok=True)
        base = folder / f"{stem}_{kind}"
        i = 1
        while base.with_suffix(".pdf").exists():
            base = folder / f"{stem}_{kind}_{i:03d}"
            i += 1
        docx_path = self._to_docx(final, base.with_suffix(".docx"))
        pdf_path = render_dover_pdf(final, base.with_suffix(".pdf"), series=series)
        self._advance_counters()
        return DoverResult(docx_path=docx_path, pdf_path=pdf_path,
                           series=series, reestr=reestr)

    def _call(self, key: str, parts: list) -> str:
        body = json.dumps({"contents": [{"parts": parts}]}).encode()
        last = ""
        for model in ("gemini-2.5-flash", "gemini-2.0-flash", "gemini-flash-latest"):
            url = ("https://generativelanguage.googleapis.com/v1beta/models/"
                   f"{model}:generateContent?key={key}")
            req = urllib.request.Request(url, data=body,
                                         headers={"Content-Type": "application/json"})
            try:
                with urllib.request.urlopen(req, timeout=120) as resp:
                    payload = json.loads(resp.read().decode())
                out = "\n".join(p.get("text", "") for p in
                                payload["candidates"][0]["content"]["parts"]).strip()
                if out:
                    return out
            except Exception as exc:  # noqa: BLE001
                last = str(exc)[:150]
        raise OfisError(f"AI javob bermadi: {last}")

    def generate(
        self,
        principal: Passport,
        agent: Passport | None,
        *,
        doc_type: str,
        description: str,
        form_date: date,
        output_dir: Path | None = None,
    ) -> DoverResult:
        text = self._compose(principal, agent, doc_type=doc_type,
                             description=description, form_date=form_date)
        stem = "".join(c if c.isalnum() or c in "_- " else "_"
                       for c in f"{principal.surname}_{principal.name}".upper()) or "DOVER"
        log.info("Dover draft for %s (%s)", principal.surname, doc_type)
        return self._save(text, stem, output_dir)

    # ------------------------------------------------------------------
    def _compose(self, principal, agent, *, doc_type, description, form_date) -> str:
        key = (self._key_getter() or "").strip()
        if not key:
            raise OfisError("AI kaliti yo'q — Sozlamalarga Gemini kalitini kiriting.")
        chosen = "" if doc_type.startswith("Авто") else f"Вид документа: {doc_type}. "
        user = (
            f"{chosen}Дата составления: {_date_dmy(form_date)}. "
            f"{_passport_block('Доверитель', principal)}. "
            f"{_passport_block('Представитель (поверенный)', agent)}. "
            f"Задание от оператора (кто, кому, для чего): {description or 'не указано'}"
        )
        body = json.dumps({
            "contents": [{"parts": [{"text": _SYSTEM + "\n\n" + user}]}],
        }).encode()
        last = ""
        for model in ("gemini-2.5-flash", "gemini-2.0-flash", "gemini-flash-latest"):
            url = ("https://generativelanguage.googleapis.com/v1beta/models/"
                   f"{model}:generateContent?key={key}")
            req = urllib.request.Request(url, data=body,
                                         headers={"Content-Type": "application/json"})
            try:
                with urllib.request.urlopen(req, timeout=90) as resp:
                    payload = json.loads(resp.read().decode())
                parts = payload["candidates"][0]["content"]["parts"]
                text = "\n".join(p.get("text", "") for p in parts).strip()
                if text:
                    return text
            except Exception as exc:  # noqa: BLE001 - try next model
                last = str(exc)[:150]
        raise OfisError(f"AI javob bermadi: {last}")

    @staticmethod
    def _to_docx(text: str, out: Path) -> Path:
        """Word copy without the blank background — same layout rules as the
        PDF (title/city/date centered, body justified, notary line split)."""
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.shared import Cm, Pt

        from src.pdf.dover_renderer import _classify

        doc = docx.Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        for role, content in _classify(text.strip().splitlines()):
            if role == "title":
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.bold = True
                run.font.size = Pt(14)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif role == "center":
                p = doc.add_paragraph(content)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif role == "notary":
                name = content.split(":", 1)[1].strip()
                p = doc.add_paragraph(f"Нотариус:\t{name}")
                p.paragraph_format.tab_stops.add_tab_stop(
                    Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
            else:
                p = doc.add_paragraph(content.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.save(str(out))
        return out
