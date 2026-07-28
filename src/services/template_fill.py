"""Filling a template the program worked out for itself.

The map comes from :mod:`src.services.template_study` and the operator has
confirmed it. What happens here is the same in all three cases — write the
value, then read the file back and prove it — but the *how* differs:

* an **AcroForm** field is set and the form flattened, so what is on screen is
  what prints;
* a **flat PDF** gets the old text redacted out of the rectangle and the new
  text typeset into it, shrinking only as far as it must (:mod:`src.pdf.rewrite`);
* a **Word** file has the value written into its paragraph or table cell, in the
  run that is already there, so it keeps the document's own font.

Every path finishes the same way: the saved file is opened again and checked
against what it was told to write. A field that did not land is named.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from src.common.errors import ValidationError
from src.common.logging import get_logger
from src.domain.documents import Passport, Patent
from src.pdf import rewrite
from src.services.template_study import DOCX, PDF_FLAT, PDF_FORM, Study

log = get_logger(__name__)


@dataclass
class FillResult:
    path: Path
    written: dict[str, str] = field(default_factory=dict)
    problems: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return not self.problems


def values_for(passport: Passport, patent: Patent | None = None, *,
               form_date: date | None = None,
               profession: str = "") -> dict[str, str]:
    """What the office knows about this worker, keyed the way the map is."""
    from src.pdf.formatters import _date_dmy

    def when(value) -> str:
        return _date_dmy(value) if value else ""

    fio = " ".join(x for x in (passport.surname, passport.name,
                               passport.patronymic or "") if x)
    out = {
        "fio": fio,
        "surname": passport.surname,
        "name": passport.name,
        "patronymic": passport.patronymic or "",
        "birth_date": when(passport.birth_date),
        "birth_place": passport.birth_place or "",
        "gender": ("Мужской" if passport.gender and passport.gender.value == "male"
                   else "Женский" if passport.gender else ""),
        "citizenship": passport.nationality or "",
        "passport_series": passport.series or "",
        "passport_number": passport.number,
        "passport_issue": when(passport.issue_date),
        "passport_issued_by": passport.issued_by or "",
        "passport_expiry": when(passport.expiry_date),
        "profession": profession or (patent.profession if patent else ""),
        "date": when(form_date or date.today()),
    }
    if patent is not None:
        out.update({"patent_series": patent.series or "",
                    "patent_number": patent.number,
                    "patent_issue": when(patent.issue_date)})
    return {k: v for k, v in out.items() if v}


# --------------------------------------------------------------------- PDF


def _fill_pdf(study: Study, template: Path, out: Path,
              values: dict[str, str]) -> FillResult:
    import fitz

    doc = fitz.open(str(template))
    result = FillResult(path=out)
    report = rewrite.Report()
    try:
        if study.kind == PDF_FORM:
            for spot in study.spots:
                value = values.get(spot.key, "")
                if not value:
                    continue
                for page in doc:
                    for widget in page.widgets() or []:
                        if widget.field_name == spot.widget:
                            widget.field_value = value
                            widget.update()
                            result.written[spot.key] = value
            # flatten, so the value is printed text and not an editable box
            doc.bake()
        else:
            for number in {s.page for s in study.spots}:
                page = doc[number - 1]
                boxes = [fitz.Rect(s.rect) for s in study.spots
                         if s.page == number and s.rect and values.get(s.key)]
                for box in boxes:
                    page.add_redact_annot(box)
                if boxes:
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

                name, font = rewrite.install_font(page)
                for spot in study.spots:
                    value = values.get(spot.key, "")
                    if spot.page != number or not spot.rect or not value:
                        continue
                    rewrite.write(page, fitz.Rect(spot.rect), value,
                                  fontname=name, font=font, size=10.0,
                                  name=spot.label, report=report)
                    result.written[spot.key] = value
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out), garbage=4, deflate=True)
    finally:
        doc.close()

    check = rewrite.verify(out, must_contain={
        spot.label: values[spot.key] for spot in study.spots
        if values.get(spot.key)})
    result.problems = check.problems() + report.problems()
    return result


# -------------------------------------------------------------------- Word


def _fill_docx(study: Study, template: Path, out: Path,
               values: dict[str, str]) -> FillResult:
    import docx

    from src.services.docx_worker import iter_paragraphs, replace_span, runs_of, text_of

    document = docx.Document(str(template))
    result = FillResult(path=out)
    paragraphs = list(iter_paragraphs(document))

    for spot in study.spots:
        value = values.get(spot.key, "")
        if not value:
            continue
        if spot.cell is not None:
            table, row, column = spot.cell
            try:
                cell = document.tables[table].rows[row].cells[column]
            except IndexError:
                result.problems.append(f"«{spot.label}» катаги топилмади")
                continue
            paragraph = cell.paragraphs[0]
            runs = runs_of(paragraph)
            text = text_of(runs)
            if runs:
                replace_span(runs, 0, len(text), value)
            else:
                paragraph.add_run(value)
        elif 0 <= spot.paragraph < len(paragraphs):
            paragraph = paragraphs[spot.paragraph]
            runs = runs_of(paragraph)
            text = text_of(runs)
            # the label stays; the value goes after it
            replace_span(runs, len(text.rstrip()), len(text), " " + value)
        else:
            result.problems.append(f"«{spot.label}» жойи топилмади")
            continue
        result.written[spot.key] = value

    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))

    written = "\n".join(p.text for p in iter_paragraphs(docx.Document(str(out))))
    for spot in study.spots:
        value = values.get(spot.key, "")
        if value and value not in written:
            result.problems.append(f"«{spot.label}» ёзилмади")
    return result


# ------------------------------------------------------------------ public


def fill(study: Study, template: Path, out: Path,
         values: dict[str, str]) -> FillResult:
    """Write the worker into the template and prove the file came out right."""
    if not template.exists():
        raise ValidationError("Шаблон топилмади", context={"path": str(template)})
    if not study.spots:
        raise ValidationError("Шаблон харитаси бўш — аввал таҳлил қилинг")

    if study.kind == DOCX:
        result = _fill_docx(study, template, out, values)
    elif study.kind in (PDF_FORM, PDF_FLAT):
        result = _fill_pdf(study, template, out, values)
    else:
        raise ValidationError("Номаълум шаблон тури", context={"kind": study.kind})

    log.info("Filled %s from its own map: %d values, %d problems",
             out.name, len(result.written), len(result.problems))
    return result
