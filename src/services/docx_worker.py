"""Finding the previous worker inside a firm's Word template and swapping them out.

Every firm words its трудовой договор differently — ГПХ, трудовой, one page or
twelve — but the worker's own data is always written the way Госуслуги write it:
«Фамилия (рус.) …», «Дата рождения …», «Серия … Номер … Дата выдачи … Кем выдан
…». So the previous worker is found by those labels rather than by position.

The label stays where it is, only the value beside it is replaced, and it is
replaced *inside the runs it already occupies* — so the value keeps the font,
size and weight the firm gave it. Everything that belongs to the firm (name,
ИНН, КПП, ОГРН, address, director, the whole legal text) is never looked at.

A value the template writes on the line *below* its label (Госуслуги do this for
«Гражданство», «Серия», «Кем выдан», «Регион», «Профессия…») is followed to the
next non-empty paragraph. A value whose old text does not look like what the
label promises is left alone rather than guessed at.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Iterable

# ------------------------------------------------------------------ shapes

_SHAPES: dict[str, re.Pattern[str]] = {
    "date": re.compile(r"^\d{1,2}\.\d{1,2}\.\d{2,4}$"),
    "gender": re.compile(r"^(мужской|женский|муж\.?|жен\.?)$", re.IGNORECASE),
    "digits": re.compile(r"^\d[\d \-]{0,20}$"),
    "code": re.compile(r"^[A-ZА-ЯЁ0-9][A-ZА-ЯЁ0-9 \-]{0,14}$", re.IGNORECASE),
    "words": re.compile(r"^[^\W\d_][\w \-\.\(\)]{0,120}$", re.UNICODE),
    "any": re.compile(r"^.{1,200}$", re.DOTALL),
}

# Labels, most specific first: the scanner takes the first alternative that
# matches at a position, so «Серия бланка» must be offered before «Серия».
_SLOTS: tuple[tuple[str, str, str], ...] = (
    ("blank_series", r"Серия\s+бланка", "code"),
    ("blank_number", r"Номер\s+бланка", "digits"),
    ("contract_date", r"Дата\s+заключения\s+договора", "date"),
    ("birth_date", r"Дата\s+рождения", "date"),
    ("pass_issue_date", r"Дата\s+выдачи", "date"),
    ("pass_issued_by", r"Кем\s+выдан(?:о)?", "any"),
    ("birth_place", r"Место\s+рождения[^,\n]*,\s*населенн\w*\s+пункт", "any"),
    ("birth_place", r"Место\s+рождения", "any"),
    ("surname", r"Фамилия\s*\(\s*рус\.?\s*\)", "words"),
    ("name", r"(?<![А-Яа-яЁё])Имя\s*\(\s*рус\.?\s*\)", "words"),
    ("patronymic", r"Отчество\s*\(\s*рус\.?\s*\)", "words"),
    ("citizenship", r"(?<![А-Яа-яЁё])Гражданство(?![А-Яа-яЁё])", "words"),
    ("gender", r"(?<![А-Яа-яЁё])Пол(?![А-Яа-яЁё])", "gender"),
    # the tab is excluded so a blank the program built keeps its «write here»
    # gap on the label's own line
    ("profession",
     r"Профессия,\s*специальность,\s*должность[^:\n\t]*:?", "any"),
    ("work_address", r"Адрес\s+места\s+работы", "any"),
    ("region", r"(?<![А-Яа-яЁё])Регион(?![А-Яа-яЁё])", "words"),
    ("series", r"(?<![А-Яа-яЁё])Серия(?![А-Яа-яЁё])", "code"),
    ("number", r"(?<![А-Яа-яЁё])Номер(?![А-Яа-яЁё])", "digits"),
    # only a line that *starts* with it — «Юридический адрес: …» is the firm's
    ("address",
     r"^\s*Адрес(?![А-Яа-яЁё])(?:\s+(?:места\s+жительства|проживания|"
     r"регистрации|пребывания))?\s*:", "any"),
)

_LABELS = re.compile("|".join(f"(?P<g{i}>{p})" for i, (_k, p, _s) in enumerate(_SLOTS)))

# «Серия»/«Номер» belong to whichever document the form was talking about last.
_PATENT_HINT = re.compile(r"разрешени\w*\s+на\s+работу|патент", re.IGNORECASE)
_PASSPORT_HINT = re.compile(r"удостоверяющ\w*\s+личност|паспорт", re.IGNORECASE)
_HINT_MAX = 90          # only a heading-length line may switch the section

#: A form's own words. A value never starts with one of these, so a line that
#: does is a heading or the next field — not the value of the label above it.
_HEADING = re.compile(
    r"^(Сведения|Документ|Вид|Выбор|Наименование|Адрес|Профессия|Статус|"
    r"Полное|Сокращ|Основной|Раздел|Работник|Работодатель|Исполнитель|\d+\.)\b")

_MONTHS = ("январ", "феврал", "март", "апрел", "ма", "июн",
           "июл", "август", "сентябр", "октябр", "ноябр", "декабр")
#: «26 мая 2026» — the «года»/«г.» after it is the firm's own wording and stays
LONG_DATE = re.compile(
    r"\d{1,2}\s+(?:" + "|".join(m + r"\w*" for m in _MONTHS) + r")\s+\d{4}",
    re.IGNORECASE)


@dataclass
class Report:
    """What the swap actually did, so the operator is never told a half-truth."""

    filled: dict[str, str] = field(default_factory=dict)
    skipped: list[str] = field(default_factory=list)
    #: ids of the paragraph elements rewritten, so a second pass leaves them be
    touched: set[int] = field(default_factory=set)

    def __bool__(self) -> bool:
        return bool(self.filled)


# ------------------------------------------------------------------- runs


def iter_paragraphs(doc) -> Iterable:
    """Every paragraph of the document, body first, then table cells."""
    yield from doc.paragraphs
    # Merged cells repeat the same underlying element; keep hard references so
    # id() values stay unique (a GC'd element's id can be reused otherwise).
    seen_ids: set[int] = set()
    keep: list = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if id(tc) in seen_ids:
                    continue
                seen_ids.add(id(tc))
                keep.append(tc)
                yield from cell.paragraphs


def runs_of(paragraph) -> list:
    """Every run in reading order — including ones inside hyperlinks.

    ``paragraph.runs`` skips those, and some converters wrap the worker's name
    in exactly that, which is how a name survives a swap unchanged.
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    return [Run(r, paragraph) for r in paragraph._p.iter(qn("w:r"))]


def text_of(runs: list) -> str:
    return "".join(r.text for r in runs)


def replace_span(runs: list, start: int, end: int, new: str) -> None:
    """Put ``new`` where characters [start, end) are, keeping their formatting.

    The first run the span touches takes the new text; the rest give up only
    the part that fell inside the span. Runs outside it are never rewritten, so
    labels, tabs and the firm's own words are untouched.

    An empty span is an insertion — the label is there but no value yet, which
    is how a blank the program built itself starts out. It goes into the last
    run that reaches the spot, so a value's own (empty) run wins over the tab
    or the label before it and the value comes out in the intended font.
    """
    if start == end:
        offset, target, at = 0, None, 0
        for run in runs:
            lo, hi = offset, offset + len(run.text)
            offset = hi
            if lo <= start <= hi:
                target, at = run, start - lo
        if target is None:
            target, at = (runs[-1], len(runs[-1].text)) if runs else (None, 0)
        if target is not None:
            target.text = target.text[:at] + new + target.text[at:]
        return

    written = False
    offset = 0
    for run in runs:
        text = run.text
        lo, hi = offset, offset + len(text)
        offset = hi
        if hi <= start or lo >= end:
            continue
        head = text[: max(0, start - lo)]
        tail = text[max(0, min(len(text), end - lo)):]
        if not written:
            run.text = head + new + tail
            written = True
        else:
            run.text = head + tail
    if not written and runs:
        runs[-1].text = runs[-1].text + new


# ------------------------------------------------------------------ scan


@dataclass(frozen=True)
class _Hit:
    key: str
    shape: str
    label_end: int
    value_start: int
    value_end: int


def _section_at(text: str, upto: int, running: str) -> str:
    """Which document the «Серия»/«Номер» at ``upto`` belongs to.

    The nearest hint *before it on the same line* wins — «Иностранный паспорт
    Серия FA Номер …» says so plainly — otherwise the last heading stands.
    """
    before = text[:upto]
    pat = max((m.end() for m in _PATENT_HINT.finditer(before)), default=-1)
    doc = max((m.end() for m in _PASSPORT_HINT.finditer(before)), default=-1)
    if pat < 0 and doc < 0:
        return running
    return "patent" if pat > doc else "passport"


def _scan(text: str, running: str) -> list[_Hit]:
    """Where each label sits and how far its value runs."""
    marks: list[tuple[str, str, int, int]] = []
    for m in _LABELS.finditer(text):
        idx = next(i for i in range(len(_SLOTS)) if m.group(f"g{i}") is not None)
        key, _pattern, shape = _SLOTS[idx]
        if key in ("series", "number"):
            section = _section_at(text, m.start(), running)
            key = ("pat_" if section == "patent" else "pass_") + key
        marks.append((key, shape, m.start(), m.end()))

    hits: list[_Hit] = []
    for i, (key, shape, _start, end) in enumerate(marks):
        limit = marks[i + 1][2] if i + 1 < len(marks) else len(text)
        lead = re.match(r"[\s:]*", text[end:limit])
        value_start = end + len(lead.group(0) if lead else "")
        value_end = value_start + len(text[value_start:limit].rstrip())
        hits.append(_Hit(key, shape, end, value_start, value_end))
    return hits


def _update_section(text: str, running: str) -> str:
    if len(text) > _HINT_MAX:
        return running
    pat = _PATENT_HINT.search(text)
    doc = _PASSPORT_HINT.search(text)
    if pat and (not doc or pat.start() < doc.start()):
        return "patent"
    if doc:
        return "passport"
    return running


def _fits(shape: str, old: str) -> bool:
    return bool(_SHAPES[shape].fullmatch(old.strip()))


# ------------------------------------------------------------------ apply


def swap_worker(doc, values: dict[str, str]) -> Report:
    """Replace the previous worker's values with ``values`` throughout ``doc``.

    ``values`` keys — surname · name · patronymic · birth_date · gender ·
    citizenship · birth_place · pass_series · pass_number · pass_issue_date ·
    pass_issued_by · pat_series · pat_number · region · blank_series ·
    blank_number · profession · contract_date · work_address · address.

    A key that is absent leaves that value as the template has it; a key mapped
    to an empty string clears it, which is how the previous worker's home
    address is stopped from travelling into the next worker's contract.
    """
    report = Report()
    section = "passport"
    lines: list[tuple[object, list, str, list[_Hit]]] = []
    for paragraph in iter_paragraphs(doc):
        runs = runs_of(paragraph)
        text = text_of(runs)
        if not text.strip():
            continue
        lines.append((paragraph, runs, text, _scan(text, section)))
        section = _update_section(text, section)

    for i, (paragraph, runs, text, hits) in enumerate(lines):
        for hit in reversed(hits):          # right to left: offsets stay valid
            if hit.key not in values:
                continue
            old = text[hit.value_start:hit.value_end]
            if old.strip():
                if not _fits(hit.shape, old):
                    report.skipped.append(hit.key)
                    continue
                replace_span(runs, hit.value_start, hit.value_end, values[hit.key])
                report.filled[hit.key] = values[hit.key]
                report.touched.add(id(paragraph._p))
                continue

            # Nothing beside the label. Either the form prints the value on the
            # line below — Госуслуги do, for «Гражданство», «Серия», «Кем
            # выдан», and then the label paragraph ends at the label — or the
            # label is followed by a gap (a tab, a space) that is asking to be
            # written on, which is how a blank the program built itself looks.
            below = (lines[i + 1]
                     if hit is hits[-1] and not text[hit.label_end:]
                     and i + 1 < len(lines) else None)
            if (below is not None and not below[3]
                    and not _HEADING.match(below[2].strip())
                    and _fits(hit.shape, below[2].strip())):
                _, next_runs, next_text, _ = below
                body = next_text.strip()
                start = next_text.index(body)
                replace_span(next_runs, start, start + len(body), values[hit.key])
                report.touched.add(id(below[0]._p))
            else:
                after = text[hit.value_start:hit.value_start + 1]
                new = values[hit.key]
                if new and after and not after.isspace():
                    new += " "
                replace_span(runs, hit.value_start, hit.value_start, new)
                report.touched.add(id(paragraph._p))
            report.filled[hit.key] = values[hit.key]

    return report


def swap_header_date(doc, new_date: str, *, limit: int = 12) -> int | None:
    """Re-date the contract's own «… 26 мая 2026 года» header line.

    Only the header is touched: a spelled-out date further down is far more
    likely to be a law or a ministry letter the firm cites, and re-dating those
    would corrupt the contract. Returns the id of the paragraph it rewrote.
    """
    for i, paragraph in enumerate(iter_paragraphs(doc)):
        if i >= limit:
            return None
        runs = runs_of(paragraph)
        text = text_of(runs)
        if _LABELS.search(text):
            return None                     # the worker block already started
        m = LONG_DATE.search(text)
        if m:
            replace_span(runs, m.start(), m.end(), new_date)
            return id(paragraph._p)
    return None
