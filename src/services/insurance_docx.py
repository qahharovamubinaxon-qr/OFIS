"""Putting a car and its drivers into any insurer's ОСАГО policy template.

Every insurer lays the policy out differently — Ингосстрах, РЕСО and Согласие
share almost no wording — but the form itself is set by the Bank of Russia, so
the *labels* are the same everywhere: «Марка, модель транспортного средства»,
«Идентификационный номер», «Государственный регистрационный знак», «Срок
страхования с … по …», and the table of «Лица, допущенные к управлению».

So the previous customer is swapped out the way the previous worker is in
:mod:`src.services.docx_worker`: by those labels, keeping each value's own font.
Some insurers print a value one character per table cell (a VIN, a date) — that
is handled too.

What this module will never do
------------------------------
* **invent a policy серия/номер.** РСА allocates those to the agency through the
  insurer; the operator records the block they were given and the program hands
  them out in order (same rule as ДМС). A made-up number is a policy that covers
  nobody.
* **touch anything that is not the car, the drivers or the dates.** The
  operator's own template is not the program's to edit. Only a span that was
  identified as the previous customer's value is replaced, in place, keeping
  its font; no line is rewritten and no line is deleted. An earlier version
  deleted the insurer's electronic-signature lines and mangled the option
  sentences — that was the program overreaching.
* **write an electronic signature.** The certificate, the МЧД and the
  доверенность belong to the insurer's signing system. They are left exactly as
  the template has them, and the operator is told they still have to come from
  the insurer.
* **invent КБМ or the premium.** Those come from the РСА database and the
  insurer's calculation; whatever the operator did not supply is left blank.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from src.common.logging import get_logger
from src.services.docx_worker import (
    replace_span,
    runs_of,
    text_of,
)

log = get_logger(__name__)


# --------------------------------------------------------------- the walk


def _all_paragraphs(doc):
    """Every paragraph of the policy, wherever Word put it.

    ``doc.paragraphs`` plus ``doc.tables`` — what the worker documents are read
    with — misses two things these policies are full of. A table nested inside
    another table's cell is one. A **text box** is the other, and one insurer
    keeps «1. Страхователь», «Собственник транспортного средства» and the
    seventeen boxes of the VIN in text boxes: read without them the собственник
    looked like a field the form did not have, and the VIN was written on a
    line of its own above boxes that still held the previous car's.

    Word stores a text box twice — once as DrawingML and once as a picture for
    older readers — and both copies are walked. They show the same thing, so
    filling only one would leave the policy saying two different things
    depending on what opened it.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    for element in doc.element.body.iter(qn("w:p")):
        yield Paragraph(element, doc)


def _all_tables(doc):
    """Every table, including the nested and the boxed ones."""
    from docx.oxml.ns import qn
    from docx.table import Table

    for element in doc.element.body.iter(qn("w:tbl")):
        yield Table(element, doc)


def _owner_of(element):
    """The cell or text box a paragraph belongs to — ``None`` for the body.

    A label may only take its value from a line in the same one of these.
    """
    from docx.oxml.ns import qn

    boundary = (qn("w:tc"), qn("w:txbxContent"))
    node = element.getparent()
    while node is not None:
        if node.tag in boundary:
            return node
        node = node.getparent()
    return None


# --------------------------------------------------------------- the labels

#: key · label · whether a colon is required before the value.
#:
#: «Страхователь» needs one. Without it the word also opens the signature line
#: «Страхователь        Представитель страховщика», and writing there replaced
#: the form's own wording with a company name.
_SLOTS: tuple[tuple[str, str, bool], ...] = (
    ("vehicle", r"Марка,?\s*модель\s+транспортного\s+средства", False),
    # «Транспортное средство:» over a column of its own — the colon is what
    # tells it apart from «2. Транспортное средство используется с прицепом»
    # and from «Паспорт транспортного средства», which are not this field.
    ("vehicle", r"Транспортное\s+средство\s*:", False),
    ("vin", r"Идентификационный\s+номер\s+транспортного\s+средства|"
            r"Идентификационный\s+номер\s*\(VIN\)|Идентификационный\s+номер", False),
    ("plate", r"Государственный\s+регистрационный\s+(?:знак|номер)"
              r"(?:\s+транспортного\s+средства)?", False),
    # Only where the form names the document itself — «Свидетельство о
    # регистрации ТС 50ое 202246». Most insurers instead head the row
    # «Паспорт транспортного средства, свидетельство о регистрации
    # транспортного средства, паспорт(свидетельство) на шасси…»: that is the
    # form's own wording for a whole block, and writing a серия over it ate
    # the sentence. Those carry «Вид документа … серия … номер …» underneath,
    # which :func:`_document_line` fills instead.
    ("sts", r"Свидетельство\s+о\s+регистрации(?:\s+ТС)?"
            r"(?!\s*(?:транспортн|ТС\b|\(|шасси))", False),
    ("policy_holder", r"(?:\d\.\s*)?Страхователь(?!\w)", True),
    # «Страхователем <ФИО> при получении настоящего страхового полиса…» on the
    # back of the policy — a sentence, so the name is always right after it
    ("policy_holder", r"(?:\d\.\s*)?Страхователем(?!\w)", False),
    ("owner", r"Собственник\s+транспортного\s+средства", True),
)
_LABELS = re.compile("|".join(f"(?P<g{i}>{p})" for i, (_k, p, _c) in enumerate(_SLOTS)))

#: Labels whose value the form prints on the line under the label rather than
#: beside it. A VIN is not among them: it has a shape of its own and belongs
#: either to its grid of boxes or to :func:`_by_pattern`.
_BELOW_KEYS = frozenset({"policy_holder", "owner", "vehicle", "plate"})
_NAME_KEYS = frozenset({"policy_holder", "owner"})

#: What the СТС is called on the «Вид документа» line once the program has
#: written one there — the previous entry is often a ПТС, and leaving that word
#: over a свидетельство's серия and номер names the wrong document.
_STS_KIND = "Свидетельство о регистрации ТС"

#: Everything after a label that is the form's own small print, not a value.
_NOTE = re.compile(r"^\s*\(?\s*(?:полное наименование|фамилия|серия|номер|"
                   r"отметить|заполняется|указывается|при наличии)", re.IGNORECASE)

#: «неограниченного количества лиц…» / «лиц, допущенных…» — the one choice the
#: operator makes that changes what the policy means.
#: The mark goes after the WHOLE option, not after the first few words of it —
#: putting it after «…к управлению» landed it in the middle of the sentence and
#: left the real box untouched.
_OPTION_TAIL = r"лиц,?\s*допущенн\w+\s+к\s+управлению\s+транспортным\s+средством\s*\d*"
_UNLIMITED = re.compile(r"неограниченн\w*\s+количеств\w*\s+" + _OPTION_TAIL,
                        re.IGNORECASE)
_LIMITED = re.compile(r"(?<!количества\s)(?<!\w)" + _OPTION_TAIL, re.IGNORECASE)
_TICKS = ("X", "Х", "[X]", "[Х]", "[x]", "[х]", "V", "✔", "✓")
_EMPTY_TICK = ("[ ]", "[]", "□", "")

#: An electronic signature belongs to the insurer, never to this program.
_SIGNATURE = re.compile(
    r"подписан\w*\s+(?:электронной|с использованием|усиленной)|"
    r"электронно[-\s]?цифровой\s+подпис|"
    r"удостоверяющий\s+центр|сертификат\s+[0-9A-F]|"
    r"мчд\s*№|доверенность\s*№|действительн\w*\s+до:", re.IGNORECASE)

_DASH = "—"


@dataclass
class Report:
    filled: dict[str, str] = field(default_factory=dict)
    drivers: int = 0
    cleared: list[str] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)


# ------------------------------------------------------------------ helpers


def _hits(text: str, values: dict[str, str]) -> list[tuple[str, int, int]]:
    """Every label on a line that this fill has a value for, with its span."""
    out: list[tuple[str, int, int]] = []
    for m in _LABELS.finditer(text):
        idx = next(i for i in range(len(_SLOTS)) if m.group(f"g{i}") is not None)
        key, _pattern, needs_colon = _SLOTS[idx]
        if key not in values:
            continue
        span = _value_span(text, m.start(), m.end(), needs_colon=needs_colon)
        if span is not None:
            out.append((key, *span))
    return out


#: A label whose value is not on this line at all — the form printed its own
#: explanation of what to write instead, and the value goes underneath.
_ELSEWHERE = (-1, -1)


def _value_span(text: str, start: int, end: int, *,
                needs_colon: bool) -> tuple[int, int] | None:
    """Where the value beside a label at [start, end) sits on the same line.

    :data:`_ELSEWHERE` when what follows is the form's own bracketed
    explanation — «1. Страхователь (полное наименование юридического лица или
    фамилия, имя, отчество физического лица)». That note is not a value; it is
    the form telling the operator to write the name on the line below, which is
    exactly where every insurer in the office's folder prints it. Reading it as
    "no such field here" is why the собственник was coming out blank.

    ``None`` when a label that demands a colon does not have one.
    """
    rest = text[end:]
    if _NOTE.match(rest):
        return _ELSEWHERE
    if needs_colon and not re.match(r"\s*:", rest):
        return None
    lead = re.match(r"[\s:№]*", rest)
    value_start = end + len(lead.group(0) if lead else "")
    stop = len(text)
    nxt = _LABELS.search(text, value_start)
    if nxt:
        stop = nxt.start()
    bracket = rest.find("(")
    if bracket >= 0 and end + bracket < stop:
        stop = end + bracket
    return value_start, value_start + len(text[value_start:stop].rstrip())


def _cells(row) -> list:
    """A table row's cells, with merged repeats collapsed."""
    return _uniq(row.cells)


def _uniq(cells) -> list:
    """The distinct cells of a list — a merged one repeats in every column.

    The cells themselves are kept, not their ids: a dropped cell's element can
    be collected and its id handed straight back out to the next one, which
    made two different columns look like the same merged cell.
    """
    out, seen = [], set()
    for cell in cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        out.append(cell)
    return out


def _set_cell(cell, value: str) -> None:
    """Write a cell, keeping the first run's formatting."""
    paragraph = cell.paragraphs[0]
    runs = runs_of(paragraph)
    text = text_of(runs)
    if runs:
        replace_span(runs, 0, len(text), value)
    else:
        paragraph.add_run(value)
    for extra in cell.paragraphs[1:]:
        for run in runs_of(extra):
            run.text = ""


def _spread(cells: list, value: str) -> bool:
    """Write one character per cell, the way a VIN grid wants it.

    Only when the row really is a grid: a run of cells each holding at most one
    character. Returns False so the caller can fall back to a plain write.
    """
    grid = [c for c in cells if len(c.text.strip()) <= 1]
    if len(grid) < 6 or len(grid) < len(value) * 0.8:
        return False
    for cell, char in zip(grid, value.ljust(len(grid))):
        _set_cell(cell, char.strip())
    return True


# ------------------------------------------------------------------- values


def _flow_lines(doc) -> list[tuple]:
    """Every line with text, as ``(owner, runs, text)``.

    ``owner`` is the cell or text box the line sits in — ``None`` for the body.
    A label may only take its value from a line with the same owner:
    «Государственный регистрационный знак транспортного средства» is the last
    line of one cell in several of these templates and the make of the car is
    the first line of the next one, so read as one flow the plate was written
    over the make.
    """
    out: list[tuple] = []
    for paragraph in _all_paragraphs(doc):
        runs = runs_of(paragraph)
        if (text := text_of(runs)).strip():
            out.append((_owner_of(paragraph._p), runs, text))
    return out


def _paragraph_values(doc, values: dict[str, str], report: Report) -> None:
    """Replace a value that sits beside its label, or on the line under it.

    Under it is the common case: the form prints «1. Страхователь (полное
    наименование юридического лица или фамилия, имя, отчество физического
    лица)» and the name goes on the next line, in its own font. Writing it
    beside the label instead put the name inside the insurer's own sentence and
    in the insurer's own colour — the blue «Марка, модель транспортного
    средства HYUNDAI SOLARIS» the operator was getting.

    A label with nothing under it either is a column *heading*; the value then
    belongs to a cell further down that only the layout knows about, and that
    is left to :func:`_by_pattern` and :func:`_column_values`.
    """
    lines = _flow_lines(doc)
    replaced: list[tuple[str, str]] = []
    for i, (cell, runs, text) in enumerate(lines):
        for key, lo, hi in reversed(_hits(text, values)):     # right to left
            old = text[lo:hi] if lo >= 0 else ""
            if old.strip():
                replace_span(runs, lo, hi, values[key])
                report.filled[key] = values[key]
                if key in _NAME_KEYS:
                    replaced.append((old.strip(), values[key]))
                continue
            if key not in _BELOW_KEYS:
                continue
            below = _below(lines, i, cell, key)
            if below is None:
                continue
            brun, btext = below
            replace_span(brun, 0, len(btext), values[key])
            report.filled[key] = values[key]
            if key in _NAME_KEYS:
                replaced.append((btext.strip(), values[key]))

    # The same name is often printed twice — once for the страхователь and once
    # for the собственник. Whatever was replaced by label is replaced elsewhere.
    for old, new in replaced:
        if len(old) < 6:
            continue
        for _p, runs, _text in lines:
            text = text_of(runs)
            at = text.find(old)
            if at >= 0 and new not in text:
                replace_span(runs, at, at + len(old), new)


#: A name or a company, not a sentence. «Оборотная сторона страхового полиса
#: обязательного страхования гражданской ответственности» was being taken for a
#: policyholder and overwritten — a heading of the insurer's own.
_NAME_MAX_CHARS = 60
_NAME_MAX_WORDS = 7


def _looks_like_a_name(text: str) -> bool:
    words = text.split()
    return bool(words) and len(text) <= _NAME_MAX_CHARS and \
        len(words) <= _NAME_MAX_WORDS


def _fits(key: str, line: str) -> bool:
    """Whether a line under a label is that label's value and nobody else's.

    «Hyundai Elantra KMHDN41BP3U633162 Т566ВЕ40» sits under the plate's heading
    in one insurer's form, but it is the whole vehicle block — writing the
    plate over it loses the make and the VIN. A line that carries somebody
    else's shape is left to :func:`_by_pattern`, which knows how to put each
    part of it back where it belongs.
    """
    return not any(other != key and shape.search(line)
                   for other, shape in _SHAPES)


def _below(lines: list, i: int, cell, key: str):
    """The line a label's value sits on, past the form's own small print.

    Only a short line inside the label's own cell qualifies, and only one that
    is not already somebody else's value: a long sentence below a label is the
    form's own text, and writing a company name over it wrecks the policy.
    """
    for entry_cell, runs, text in lines[i + 1:i + 4]:
        if entry_cell is not cell:
            return None                # the next cell, not this label's value
        if _NOTE.match(text.strip()) or not text.strip():
            continue
        if _LABELS.search(text):
            return None                # the next field, not this one's value
        stripped = text.strip()
        if not _looks_like_a_name(stripped) or not _fits(key, stripped):
            return None
        return runs, text
    return None


def _heading(cell, values: dict[str, str], skip: frozenset[str]) -> str | None:
    """The field a cell is the *heading* of — a label and nothing else.

    ``None`` when the cell already carries its value, when the label is only
    part of a longer sentence of the insurer's, or when the value has already
    been written somewhere better.
    """
    text = " ".join(cell.text.split())
    if not text:
        return None
    found = _hits(text, values)
    if len(found) != 1:
        return None
    key, lo, hi = found[0]
    if lo < 0 or text[lo:hi].strip() or hi < len(text.rstrip()):
        return None
    return None if key in skip else key


def _column_values(doc, values: dict[str, str], report: Report, *,
                   skip: frozenset[str]) -> None:
    """A heading over a column, with the value in the row underneath.

    «Транспортное средство: | Идентификационный номер | Государственный
    регистрационный знак» across one row and «KIA | KNEDE22126611237 |
    X526AY550» spread over the next two is how one of the office's insurers
    lays the car out, and the make and model get a row each. Merged cells make
    the columns of two rows line up by grid position rather than by count, so
    that is what is followed.
    """
    # what is already placed is decided once: Word keeps two copies of every
    # text box and both have to end up saying the same thing
    done = skip | frozenset(report.filled)
    for table in _all_tables(doc):
        grid = [list(row.cells) for row in table.rows]
        for ri, row_cells in enumerate(grid):
            for cell in _uniq(row_cells):
                key = _heading(cell, values, done)
                if key is None:
                    continue
                columns = [i for i, c in enumerate(row_cells)
                           if c._tc is cell._tc]
                targets = _value_cells(grid, ri, columns, cell, key)
                if not targets:
                    continue
                _stack(targets, values[key])
                report.filled[key] = values[key]


def _value_cells(grid: list, ri: int, columns: list[int], heading,
                 key: str) -> list:
    """The cells under a heading that hold the previous car's value.

    Only a short, value-shaped cell qualifies. The row under a heading is as
    often the insurer's own paragraph spanning the whole table — «Паспорт
    транспортного средства, свидетельство о регистрации транспортного
    средства…» — and writing a VIN over that deletes the form.
    """
    written: list = []
    empty: list = []
    for row_cells in grid[ri + 1:ri + 3]:
        picked = _uniq([row_cells[i] for i in columns if i < len(row_cells)])
        for cell in picked:
            if cell._tc is heading._tc:
                continue                      # the heading, merged downwards
            text = " ".join(cell.text.split())
            if not text:
                empty.append(cell)
                continue
            if _LABELS.search(text) or not _looks_like_a_name(text) \
                    or not _fits(key, text):
                continue
            written.append(cell)
    return written or empty[:1]


def _stack(cells: list, value: str) -> None:
    """Write a value over the cells that held the previous one.

    Two cells for «KIA» and «RIO» take a word each, so the new car keeps the
    make-over-model shape the template was built with; anything left over goes
    into the last of them, and a spare cell is cleared rather than left showing
    half of the last customer's car.
    """
    if len(cells) == 1:
        _set_cell(cells[0], value)
        return
    words = value.split()
    if len(words) >= len(cells):
        parts = words[:len(cells) - 1] + [" ".join(words[len(cells) - 1:])]
    else:
        parts = words + [""] * (len(cells) - len(words))
    for cell, part in zip(cells, parts, strict=True):
        _set_cell(cell, part)


def _cell_gaps(doc, values: dict[str, str], report: Report, *,
               skip: frozenset[str]) -> None:
    """A labelled cell the template left blank — give the value its own line.

    One insurer prints «Идентификационный номер транспортного средства» and
    nothing under it while the cells either side of it carry «Марка, модель…»
    over the make and «Государственный регистрационный знак…» over the plate.
    The value goes on a line of its own copied from one of those neighbours, so
    it comes out in the black the template writes values in — appending it to
    the label put the VIN inside the insurer's blue heading.
    """
    done = skip | frozenset(report.filled)
    for table in _all_tables(doc):
        for row in table.rows:
            cells = _cells(row)
            donor = next((p for c in cells for p in _value_lines(c)), None)
            if donor is None:
                continue
            for cell in cells:
                key = _heading(cell, values, done)
                if key is None:
                    continue
                _add_line(cell, donor, values[key])
                report.filled[key] = values[key]


def _value_lines(cell) -> list:
    """The paragraphs of a cell that hold a value rather than a label."""
    out = []
    for paragraph in cell.paragraphs[1:]:
        runs = runs_of(paragraph)
        text = text_of(runs)
        if text.strip() and not _LABELS.search(text) and not _NOTE.match(text.strip()):
            out.append(paragraph)
    return out


def _add_line(cell, donor, value: str) -> None:
    """Put ``value`` on a new line of ``cell``, in ``donor``'s formatting."""
    import copy

    from docx.text.paragraph import Paragraph

    element = copy.deepcopy(donor._p)
    cell._tc.append(element)
    runs = runs_of(Paragraph(element, cell))
    if runs:
        replace_span(runs, 0, len(text_of(runs)), value)


# ------------------------------------------------- «Вид документа … серия …»

_DOC_KIND = re.compile(r"Вид\s+документа", re.IGNORECASE)
_DOC_SERIES = re.compile(r"сери[яи]\s*[:№]?\s*(\S+)", re.IGNORECASE)
_DOC_NUMBER = re.compile(r"номер\s*[:№]?\s*(\S+)", re.IGNORECASE)


def _document_line(doc, sts: str, report: Report) -> None:
    """Fill «Вид документа … серия … номер …» with the car's СТС.

    Most of the office's insurers head the row with the whole of the form's own
    wording — «Паспорт транспортного средства, свидетельство о регистрации
    транспортного средства, паспорт(свидетельство) на шасси…» — and print the
    document itself on the line under it. The kind is rewritten too: the line
    is as likely to say «Паспорт ТС» from the previous car, and a свидетельство
    filed under a паспорт's name is the wrong document.

    A line with no digits in it is the form's own «Вид документа Серия Номер»
    column heading, and is left alone.
    """
    if not sts:
        return
    words = sts.split()
    series = " ".join(words[:-1]) if len(words) > 1 else ""
    number = words[-1] if words else ""
    for paragraph in _all_paragraphs(doc):
        runs = runs_of(paragraph)
        text = text_of(runs)
        head = _DOC_KIND.search(text)
        if head is None or not any(c.isdigit() for c in text):
            continue
        found_series = _DOC_SERIES.search(text, head.end())
        found_number = _DOC_NUMBER.search(text, head.end())
        edits: list[tuple[int, int, str]] = []
        if found_series is not None and series:
            edits.append((*found_series.span(1), series))
        if found_number is not None and number:
            edits.append((*found_number.span(1), number))
        stop = min([m.start() for m in (found_series, found_number)
                    if m is not None] or [len(text)])
        kind = text[head.end():stop]
        named = kind.strip(" \t:№.,")
        if named:
            at = head.end() + kind.index(named)
            edits.append((at, at + len(named),
                          _STS_KIND.upper() if named.isupper() else _STS_KIND))
        for lo, hi, value in sorted(edits, reverse=True):     # right to left
            replace_span(runs, lo, hi, value)
        if edits:
            report.filled["sts"] = sts


# --------------------------------------------------- by shape, not by label

#: The previous customer's data, recognised by how it is written. A VIN is 17
#: characters with no I/O/Q; a Russian plate is a letter, three digits, two
#: letters and a region; a СТС is two digits, two letters and six digits.
#: A plate's letters are the twelve that look the same in both alphabets, and
#: templates are typed in whichever the keyboard was on — so both are accepted.
_PLATE_LETTER = "АВЕКМНОРСТУХABEKMHOPCTYX"
_SHAPES: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("vin", re.compile(r"(?<![A-Z0-9])(?=[A-Z0-9]*\d)(?=[A-Z0-9]*[A-Z])"
                       r"[A-HJ-NPR-Z0-9]{15,17}(?![A-Z0-9])")),
    ("plate", re.compile(rf"(?<![\w])[{_PLATE_LETTER}]\s?\d{{3}}\s?"
                         rf"[{_PLATE_LETTER}]{{2}}\s?\d{{2,3}}(?!\w)")),
    ("sts", re.compile(r"(?<!\d)\d{2}\s?[А-Яа-я]{2}\s?\d{6}(?!\d)")),
)


def _by_pattern(doc, values: dict[str, str], report: Report) -> None:
    """Swap the previous customer's car for this one wherever it is written.

    Layout-independent: an insurer that prints its headings on one line and the
    data on another is handled here rather than by guessing which line is which.
    Make and model have no shape of their own, so they are taken to be whatever
    stands in front of the VIN on the same line.
    """
    for paragraph in _all_paragraphs(doc):
        _swap_shapes(runs_of(paragraph), values, report)


_VEHICLE_HEAD = re.compile(r"^\s*([A-Za-zА-Яа-я][\w\-. ]{2,40}?)\s+$")


def _swap_shapes(runs: list, values: dict[str, str], report: Report) -> None:
    text = text_of(runs)
    if not text.strip():
        return
    for key, shape in _SHAPES:
        new = values.get(key)
        if not new:
            continue
        for m in reversed(list(shape.finditer(text))):
            if m.group(0).replace(" ", "") == new.replace(" ", ""):
                continue                       # already this car
            if key == "vin" and values.get("vehicle"):
                head = _VEHICLE_HEAD.match(text[:m.start()])
                if head and not _LABELS.search(head.group(1)):
                    replace_span(runs, head.start(1), head.end(1),
                                 values["vehicle"])
                    report.filled["vehicle"] = values["vehicle"]
                    text = text_of(runs)
                    m = next((x for x in shape.finditer(text)), m)
            replace_span(runs, m.start(), m.end(), new)
            report.filled[key] = new
            text = text_of(runs)


_VIN = re.compile(r"^[A-HJ-NPR-Z0-9]{15,17}$")
_VIN_CHARS = re.compile(r"^[A-HJ-NPR-Z0-9]*$")
_VIN_LABEL = re.compile(r"Идентификационный\s+номер")


def _boxes(cells: list) -> list:
    """The longest unbroken run of cells holding at most one character each."""
    best: list = []
    run: list = []
    for cell in cells:
        if len(cell.text.strip()) <= 1:
            run.append(cell)
            best = run if len(run) > len(best) else best
        else:
            run = []
    return best


def _vin_grids(doc, vin: str, report: Report) -> None:
    """Some insurers print the VIN one character per box — rewrite those too.

    A row of fifteen boxes or more is the VIN's: a date takes eight and no
    other field on the form is that long. The boxes are taken as they come,
    whether the previous car's VIN is still standing in them, only part of it
    is, or the template ships them empty — requiring them to spell out a whole
    valid VIN left half-filled grids untouched and the VIN printed above them
    instead of inside them.
    """
    if not vin:
        return
    for table in _all_tables(doc):
        rows = list(table.rows)
        for ri, row in enumerate(rows):
            cells = _cells(row)
            boxes = _boxes(cells)
            if len(boxes) < 15:
                continue
            joined = "".join(c.text.strip() for c in boxes).upper()
            if not _VIN.fullmatch(joined) and not (
                    _VIN_CHARS.fullmatch(joined) and _labelled_vin(rows, ri, cells)):
                continue
            for cell, char in zip(boxes, vin.ljust(len(boxes))):
                _set_cell(cell, char.strip())
            report.filled["vin_grid"] = vin


def _labelled_vin(rows: list, ri: int, cells: list) -> bool:
    """Whether «Идентификационный номер» heads this row of boxes."""
    near = [c.text for c in cells]
    for row in rows[max(0, ri - 2):ri]:
        near += [c.text for c in _cells(row)]
    return any(_VIN_LABEL.search(text) for text in near)


# ------------------------------------------------------------------ drivers

_DRIVER_HEADER = re.compile(
    r"Лица,\s*допущенн\w+\s+к\s+управлению\s+транспортным\s+средством",
    re.IGNORECASE)
_LICENCE_HEADER = re.compile(r"Водительское\s+удостоверение", re.IGNORECASE)


def _driver_table(doc):
    """The table whose header row names the drivers and their licences."""
    for table in _all_tables(doc):
        for ri, row in enumerate(table.rows):
            joined = " ".join(c.text for c in _cells(row))
            if _DRIVER_HEADER.search(joined) and _LICENCE_HEADER.search(joined):
                return table, ri
    return None, 0


def _is_driver_row(cells: list) -> bool:
    """A row that holds a driver — not a footnote or a spacer under the table."""
    if len(cells) < 3:
        return False
    texts = [c.text.strip() for c in cells]
    if texts[0].isdigit():
        return True
    return sum(1 for t in texts[1:] if t) >= 2


def _column_numbers(cells: list) -> bool:
    """«1 | 2 | 3 | 4» — the form numbering its columns, not a driver row.

    Its first cell is «1» just like a real first driver's, so the whole row has
    to be read: consecutive integers all the way across and nothing else.
    """
    texts = [c.text.strip() for c in cells]
    if len(texts) < 3 or not all(t.isdigit() for t in texts):
        return False
    return [int(t) for t in texts] == list(range(1, len(texts) + 1))


def _fill_drivers(doc, drivers: list[tuple[str, str]], report: Report) -> None:
    """Write each driver into the table, and dash out the rows left over.

    The КБМ column is not touched: that coefficient comes from the РСА database
    and this program has no way to know it.
    """
    table, header = _driver_table(doc)
    if table is None:
        if not _fill_driver_lines(doc, drivers, report):
            report.skipped.append("drivers — бланкада ҳайдовчилар рўйхати топилмади")
        return
    body = [row for row in table.rows[header + 1:] if _cells(row)]
    while body and _column_numbers(_cells(body[0])):
        body = body[1:]
    body = [row for row in body if _is_driver_row(_cells(row))]
    for i, row in enumerate(body):
        cells = _cells(row)
        if len(cells) < 3:
            continue
        fio, licence = drivers[i] if i < len(drivers) else (_DASH, _DASH)
        _set_cell(cells[1], fio)
        _set_cell(cells[2], licence)
        if i >= len(drivers) and len(cells) > 3:
            _set_cell(cells[3], _DASH)
    report.drivers = min(len(drivers), len(body))
    if len(drivers) > len(body):
        report.skipped.append(
            f"drivers:{len(drivers) - len(body)} — бланкада фақат "
            f"{len(body)} та қатор бор")


#: «ЧУНАЕВ БАХРОМЖОН ЭШМИРЗОЕВИЧ 5036 634917 0.63» — a driver written as a
#: plain line, which is how an insurer that draws no table prints them.
_DRIVER_LINE = re.compile(
    r"^(?P<fio>[А-ЯЁ][А-ЯЁа-яё\-]+(?:\s+[А-ЯЁ][А-ЯЁа-яё\-]+){1,3})\s+"
    r"(?P<licence>[A-ZА-Я0-9]{2,4}\s?\d{6,9})"
    r"(?P<tail>\s+\d[.,]\d{1,2})?\s*$")


def _fill_driver_lines(doc, drivers: list[tuple[str, str]],
                       report: Report) -> bool:
    """Rewrite a driver list an insurer printed as paragraphs, not a table.

    The КБМ that trails each line belongs to the person who was there before,
    so it is dropped rather than handed to somebody else — it comes from the
    РСА database, and this program has no way to look it up.
    """
    lines = [(p, r, text_of(r)) for p in _all_paragraphs(doc)
             if (r := runs_of(p)) and text_of(r).strip()]
    start = next((i for i, (_p, _r, t) in enumerate(lines)
                  if _DRIVER_HEADER.search(t)), None)
    if start is None:
        return False
    block: list[tuple] = []
    for entry in lines[start + 1:]:
        if _DRIVER_LINE.match(entry[2].strip()):
            block.append(entry)
        elif block:
            break
    if not block:
        return False
    for i, (_p, runs, text) in enumerate(block):
        if i < len(drivers):
            fio, licence = drivers[i]
            new = f"{fio} {licence}"
        else:
            new = f"{_DASH} {_DASH}"
        replace_span(runs, 0, len(text), new)
    report.drivers = min(len(drivers), len(block))
    if len(drivers) > len(block):
        report.skipped.append(
            f"drivers:{len(drivers) - len(block)} — бланкада фақат "
            f"{len(block)} та қатор бор")
    return True


# -------------------------------------------------------------- the choice


_MARK = re.compile(r"\s*(\[[^\]]{0,3}\]|[XХxхV✔✓])")


def _tick(text: str, m: re.Match, on: bool) -> tuple[int, int, str] | None:
    """The span of the mark beside one option, and what it should become.

    Only the mark itself — never the sentence around it. Rewriting the whole
    line was what mangled the operator's own template.
    """
    marked = _MARK.match(text, m.end())
    if on:
        if marked and marked.group(1).strip("[] ").strip():
            return None                       # already ticked
        if marked:
            return marked.start(1), marked.end(1), "[X]"
        return m.end(), m.end(), " [X]"
    if not marked or not marked.group(1).strip("[] ").strip():
        return None                           # nothing to clear
    empty = "[ ]" if marked.group(1).startswith("[") else ""
    return marked.start(1), marked.end(1), empty


def _choose(doc, unlimited: bool, report: Report) -> None:
    """Mark «неограниченного количества лиц» or «лиц, допущенных к управлению».

    Which line is ticked decides who the policy covers, and it follows from the
    photographs the operator uploaded. Only the mark is written — everything
    the insurer printed stays exactly as it is.
    """
    wanted = _UNLIMITED if unlimited else _LIMITED
    offered = False
    for paragraph in _all_paragraphs(doc):
        runs = runs_of(paragraph)
        text = text_of(runs)
        if not (_UNLIMITED.search(text) or _LIMITED.search(text)):
            continue
        offered = offered or bool(wanted.search(text))
        edits = []
        for pattern, on in ((_UNLIMITED, unlimited), (_LIMITED, not unlimited)):
            m = pattern.search(text)
            if m:
                edit = _tick(text, m, on)
                if edit:
                    edits.append(edit)
        for lo, hi, mark in sorted(edits, reverse=True):   # right to left
            replace_span(runs, lo, hi, mark)
    report.filled["coverage"] = ("неограниченное количество лиц" if unlimited
                                 else "лица, допущенные к управлению")
    if not offered:
        # this insurer's template prints only the other option — ticking it
        # would say something the form does not offer
        report.skipped.append("coverage — бланкада бу вариант йўқ")


# ----------------------------------------------------------- the signature


def _note_signature(doc, report: Report) -> None:
    """Notice the insurer's electronic signature — and leave it alone.

    An earlier version deleted these lines, reasoning that a certificate from
    the previous policy should not travel onto a new one. That was the program
    destroying the operator's own template, which is not its call to make.
    The template is left exactly as it was given, and the operator is told what
    still has to come from the insurer.
    """
    for paragraph in _all_paragraphs(doc):
        text = text_of(runs_of(paragraph))
        if text.strip() and _SIGNATURE.search(text):
            report.cleared.append(" ".join(text.split())[:60])
            return


# ------------------------------------------------- a line, however it is drawn
#
# The Bank of Russia's form asks for the срок twice — «Срок страхования с … по
# …» and «Страхование распространяется … с … по …» — and leaves it to the
# insurer how to print the eight digits of each date. Three ways turn up in the
# office's folder:
#
#   * as text, «15.07.2026», which needs nothing clever;
#   * one digit per table cell, interleaved with cells holding the form's own
#     «.» and «.20» — «с | 2 | 8 | . | 0 | 6 | .20 | 2 | 4 | г.»;
#   * one digit per little text box, drawn side by side in a group, in whatever
#     order they happened to be created.
#
# A fourth turns up inside the third: four boxes floating *over* a line, which
# Word anchors at the start of the paragraph they sit in the middle of.
#
# All four are read here as one thing — a *strip*: the characters of one
# printed line, in the order the form is read, each remembering the run it came
# from. A date found in a strip can then be written back digit over digit,
# which is the whole point: the tabs, the dots and the «20» the template
# printed itself must not move, or the digits stop landing in their boxes.

#: One piece of a printed line: the run lists that spell it, and the slice of
#: their text it contributes. There is more than one run list when Word keeps
#: several copies of the same box — the one it draws and the picture older
#: readers fall back on — and every copy has to end up saying the same thing.
_Piece = tuple[list[list], int, int]

_DML = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_VML = "{urn:schemas-microsoft-com:vml}"
_WPG = "{http://schemas.microsoft.com/office/word/2010/wordprocessingGroup}"
_WPS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}"
_GROUP_TAGS = (_WPG + "wgp", _VML + "group")
_SHAPE_TAGS = (_WPS + "wsp", _VML + "shape", _VML + "rect", _VML + "roundrect")


def _own_runs(paragraph) -> list:
    """A paragraph's own runs — not those of a box floating over it.

    :func:`~src.services.docx_worker.runs_of` walks everything underneath a
    paragraph, and a text box anchored in it hangs there too. On one insurer's
    form the four boxes holding the end date's day and month are anchored at
    the *start* of the line they are drawn in the middle of, so read that way
    the line begins «2706». They are read where they are drawn instead — see
    :func:`_spliced`.
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    floating = (qn("w:drawing"), qn("w:pict"))
    out = []
    for element in paragraph._p.iter(qn("w:r")):
        node = element.getparent()
        while node is not None and node is not paragraph._p:
            if node.tag in floating:
                break
            node = node.getparent()
        else:
            out.append(Run(element, paragraph))
    return out


def _paragraph_pieces(paragraph) -> list[_Piece]:
    runs = _own_runs(paragraph)
    return [([runs], 0, len(text_of(runs)))] if runs else []


def _row_pieces(row) -> list[_Piece]:
    """A table row read across, cell by cell.

    This is the row of boxes: the date is spelled by cells that hold one
    character each, and the cells between them hold the form's own «.» and
    «.20». Read cell by cell it comes back out as «28.06.2024».
    """
    return [piece for cell in _cells(row)
            for paragraph in cell.paragraphs
            for piece in _paragraph_pieces(paragraph)]


def _strip_text(pieces: list[_Piece]) -> tuple[str, list[tuple[int, int]], list[bool]]:
    """The pieces read as one line, with the whitespace taken out.

    ``(text, where, fresh)`` — ``where[i]`` is the piece and offset character
    ``i`` came from, and ``fresh[i]`` whether a space, a tab or a new piece
    stood in front of it. The whitespace has to go, or «2 8 . 0 6 . 20 2 4»
    never looks like a date; ``fresh`` is what is left of it, and it is what
    tells «в течение срока страхования с 15.07.2026» from a word that merely
    ends in «с».
    """
    text: list[str] = []
    where: list[tuple[int, int]] = []
    fresh: list[bool] = []
    for index, (runs, lo, hi) in enumerate(pieces):
        gap = True
        for offset, char in enumerate(text_of(runs[0])[lo:hi], start=lo):
            if char.isspace():
                gap = True
                continue
            text.append(char)
            where.append((index, offset))
            fresh.append(gap)
            gap = False
    return "".join(text), where, fresh


def _strip_raw(pieces: list[_Piece]) -> str:
    """The same line with its wording intact — for reading it, never writing."""
    return " ".join(part for runs, lo, hi in pieces
                    if (part := " ".join(text_of(runs[0])[lo:hi].split())))


def _put(pieces: list[_Piece], where: list[tuple[int, int]],
         lo: int, hi: int, new: str) -> None:
    """Write ``new`` over characters [lo, hi) of a strip — one for one.

    Never more and never fewer characters, so every tab, dot and «20» the
    template printed itself stays exactly where it was and each digit keeps
    landing in its own box.
    """
    for index, char in zip(range(lo, hi), new, strict=True):
        piece, offset = where[index]
        for runs in pieces[piece][0]:
            replace_span(runs, offset, offset + 1, char)


# ---------------------------------------------------------- boxes in a group


def _shape_box(shape) -> tuple[int, int, int, int] | None:
    """Where inside its group a shape is drawn — (left, top, width, height).

    Word writes the same group twice: in DrawingML, where the numbers are EMU
    on ``<a:off>``/``<a:ext>``, and as a VML picture, where they are group
    units in the ``style`` attribute. Only the order they give matters here,
    and both orders are the same.
    """
    for xfrm in shape.iter(_DML + "xfrm"):
        off, ext = xfrm.find(_DML + "off"), xfrm.find(_DML + "ext")
        if off is None or ext is None:
            return None
        return (int(off.get("x")), int(off.get("y")),
                int(ext.get("cx")), int(ext.get("cy")))
    style = dict(part.split(":", 1) for part in
                 (shape.get("style") or "").split(";") if ":" in part)
    try:
        return (round(float(style["left"])), round(float(style["top"])),
                round(float(style["width"])), round(float(style["height"])))
    except (KeyError, ValueError):
        return None


def _band_strips(group) -> list[list[_Piece]]:
    """One group of little text boxes, read as the lines it draws.

    «Срок страхования с», the hour, the minute, the day, the month, the year:
    each is its own box, and they are stored in the order they were drawn,
    which is not the order they are read. Put back down the page and then
    across, they spell the row again.

    A separator the form draws once for both rows — the «.» between day and
    month, the «. 20» before the year — is one box tall enough to reach both,
    holding a paragraph for each: its first line belongs to the first row and
    its second to the second. Reading the whole box into both rows instead
    gives «28..03.20.2026», which is no date at all.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    shapes: list[tuple[tuple[int, int, int, int], list[list[_Piece]]]] = []
    seen: set[tuple] = set()
    for shape in group.iter(*_SHAPE_TAGS):
        box = _shape_box(shape)
        if box is None:
            continue
        lines = [pieces for content in shape.iter(qn("w:txbxContent"))
                 for element in content.iter(qn("w:p"))
                 if (pieces := _paragraph_pieces(Paragraph(element, None)))]
        # the same box drawn twice over is one box, not two
        if not lines or (key := (box, tuple(map(_strip_raw, lines)))) in seen:
            continue
        seen.add(key)
        shapes.append((box, lines))
    if not shapes:
        return []

    height = min(box[3] for box, _lines in shapes)     # the shortest box is one row
    tops: list[int] = []
    for top in sorted({box[1] for box, _lines in shapes}):
        if not tops or top - tops[-1] > height // 2:
            tops.append(top)

    bands: list[list[tuple[int, list[_Piece]]]] = [[] for _top in tops]
    for box, lines in shapes:
        spans = [i for i, top in enumerate(tops)
                 if box[1] < top + height and box[1] + box[3] > top]
        if not spans:
            continue
        # a box as tall as the rows it reaches has a line for each of them;
        # anything else is one box on its own row
        share = (list(zip(spans, lines, strict=True)) if len(lines) == len(spans)
                 else [(spans[0], [piece for line in lines for piece in line])])
        for band, pieces in share:
            bands[band].append((box[0], pieces))
    return [[piece for _left, pieces in sorted(row, key=lambda item: item[0])
             for piece in pieces] for row in bands]


# ------------------------------------------------- boxes floating over a line


#: A cover date whose day and month were left to boxes floating over the line —
#: «по . .20 2 5» — with a year still printed, so there is something to check
#: the boxes against. The blank extra periods the form offers underneath («с .
#: . 20 г. по . .20 г.») have no year and are left alone.
_HOLLOW = re.compile(r"(?:с|c|по)\.\.(?=\d{4}(?!\d))", re.IGNORECASE)

#: How many digits such a hole wants back: a day and a month.
_HOLLOW_DIGITS = 4


def _paragraph_element(runs: list):
    from docx.oxml.ns import qn

    node = runs[0]._r if runs else None
    while node is not None and node.tag != qn("w:p"):
        node = node.getparent()
    return node


def _floating_boxes(pieces: list[_Piece]) -> list[list[list]]:
    """Rows of single-character boxes drawn over a line, one list per copy.

    Empty when there are none, and empty too when the copies do not all hold
    the same thing: that is a template already saying two things, and picking
    one of them would be a guess.
    """
    from docx.oxml.ns import qn
    from docx.table import Table

    grids: list[list[list]] = []
    spelled: set[str] = set()
    seen: set[int] = set()
    for runs, _lo, _hi in pieces:
        element = _paragraph_element(runs[0])
        if element is None or id(element) in seen:
            continue
        seen.add(id(element))
        for tbl in element.iter(qn("w:tbl")):
            for row in Table(tbl, None).rows:
                cells = _boxes(_cells(row))
                joined = "".join(cell.text.strip() for cell in cells)
                if len(cells) != _HOLLOW_DIGITS or not joined.isdigit():
                    continue
                grids.append([runs_of(cell.paragraphs[0]) for cell in cells])
                spelled.add(joined)
    return grids if len(spelled) == 1 else []


def _spliced(pieces: list[_Piece]) -> list[_Piece] | None:
    """A line whose day and month float over it in boxes, read whole.

    One insurer prints the end of «Страхование распространяется … по …» as four
    little boxes drawn on top of the line. The line itself reads «по . .20 2
    5»: the day and the month are simply not in it, and the boxes that hold
    them are anchored, as Word anchors them, at the start of the paragraph.
    Read with the boxes back in their holes it says «по 27.06.2025» again —
    a date that can be recognised, checked and written over.

    ``None`` when there is no such hole, when the boxes are not there, or when
    the two dots are not on one piece: filling a hole on a guess would leave
    the policy carrying a date that was never anybody's cover.
    """
    text, where, _fresh = _strip_text(pieces)
    hole = _HOLLOW.search(text)
    if hole is None:
        return None
    grids = _floating_boxes(pieces)
    if not grids:
        return None
    (index, first), (again, second) = where[hole.end() - 2], where[hole.end() - 1]
    if index != again:
        return None
    runs, lo, hi = pieces[index]
    day, month = ([[grid[i] for grid in grids] for i in range(2)],
                  [[grid[i] for grid in grids] for i in (2, 3)])
    return (pieces[:index]
            + [(runs, lo, first)] + [(copies, 0, 1) for copies in day]
            + [(runs, first, second)] + [(copies, 0, 1) for copies in month]
            + [(runs, second, hi)]
            + pieces[index + 1:])


# ------------------------------------------------------- which date is a срок

#: «14.07.2027г.» — the «г.» runs straight into the year on some forms,
#: so a word boundary would miss it.
_DMY = re.compile(r"(?<!\d)\d{2}\.\d{2}\.\d{4}(?!\d)")

#: A date the form itself introduces with «с …» or «по …», through the insurer's
#: own «00 ч. 01 мин.» where there is one. That is what tells the cover apart
#: from «Дата заключения договора 28.06.2024 г.» and «Дата выдачи полиса», which
#: carry the same date and are not the срок.
_COVER = re.compile(r"(?:с|c|по)(?:\d{1,2}ч\.\d{1,2}мин\.)?"
                    rf"({_DMY.pattern})", re.IGNORECASE)

#: «Действителен с 08.12.2025 по 07.05.2027» under an insurer's certificate is
#: how long their signing key lasts, not how long the car is covered.
_VALIDITY = re.compile(r"Действител[а-яё]*(?:с|до)", re.IGNORECASE)


def _cover_on(text: str, fresh: list[bool]) -> list[tuple[int, int, str]]:
    """The cover dates in a strip's squashed text, as (start, end, value)."""
    if _VALIDITY.search(text):
        return []
    return [(m.start(1), m.end(1), m.group(1)) for m in _COVER.finditer(text)
            # a word that merely ends in «с» is not the form saying «с»
            if fresh[m.start()]]


def _cover_dates(pieces: list[_Piece]) -> list[tuple[int, int, str]]:
    """The cover dates on one line, as (start, end, value) in its squashed text."""
    text, _where, fresh = _strip_text(pieces)
    return _cover_on(text, fresh)


# --------------------------------------------------------------- every line


def _strips(doc):
    """Every printed line of the policy, in the order the form is read."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for element in doc.element.body.iter(qn("w:p"), qn("w:tbl"), *_GROUP_TAGS):
        if element.tag == qn("w:p"):
            found = [_paragraph_pieces(Paragraph(element, doc))]
        elif element.tag == qn("w:tbl"):
            found = [_row_pieces(row) for row in Table(element, doc).rows]
        else:
            found = _band_strips(element)
        for pieces in found:
            if not pieces:
                continue
            yield pieces
            if (spliced := _spliced(pieces)) is not None:
                yield spliced


# ------------------------------------------------------------------ public


#: Capitalised: the same words appear mid-sentence in the small print.
_TERM = re.compile(r"(?<![А-Яа-яЁё])Срок\s+страхования")
#: The form asks for the same срок a second time, in a row of its own — and
#: that is the row the office kept getting back with the previous cover on it.
_SPREAD = re.compile(r"Страхование\s+распространяется")
_SIGNED_ON = re.compile(r"Дата\s+(?:заключения\s+договора|выдачи\s+полиса)",
                        re.IGNORECASE)
_LONG_DATE = re.compile(
    r"«?\d{1,2}»?\s+[А-Яа-яЁё]{3,8}\s+\d{4}", re.IGNORECASE)


def _runs_forward(start: str, end: str) -> bool:
    """Whether a pair of dates could be a cover — the end after the start."""
    def order(dmy: str) -> tuple[int, ...]:
        day, month, year = dmy.split(".")
        return int(year), int(month), int(day)

    return order(start) < order(end)


def _names_the_term(pieces: list[_Piece]) -> bool:
    """Whether a line is one of the two the form asks for the срок on."""
    wording = _strip_raw(pieces)
    return bool(_TERM.search(wording) or _SPREAD.search(wording))


def _previous_cover(lines: list[list[_Piece]]) -> tuple[str, str] | None:
    """What the blank's own cover ran — its start, and its end.

    The first two different dates the form introduces with «с …» and «по …»,
    once it has named the срок. Reading them off the blank is what lets every
    other copy of them be found by value rather than by position, and a form
    prints the pair in more places than any one rule would find: twice over on
    the face of it, and twice again in the copy Word keeps of every text box.

    ``None`` when the pair does not run forwards, or when there is no pair —
    a blank whose boxes were never filled in has nothing to match, and is
    handed to the operator rather than guessed at.
    """
    seen: list[str] = []
    named = False
    for pieces in lines:
        if not named and not (named := _names_the_term(pieces)):
            continue
        for _lo, _hi, value in _cover_dates(pieces):
            if value in seen:
                continue
            seen.append(value)
            if len(seen) == 2:
                return (seen[0], seen[1]) if _runs_forward(*seen) else None
    return None


def _cover(doc, *, start: str, end: str, report: Report) -> None:
    """Put this policy's cover wherever the blank printed the one before it.

    The form asks for the срок twice — «Срок страхования с … по …» and
    «Страхование распространяется … с … по …» — and leaves each insurer to
    print the eight digits of a date however it likes: as text, one digit per
    table cell, one digit per little text box, or with four of the digits in
    boxes floating over the line. :func:`_strips` reads all of them back as the
    line they spell, and the new date goes in digit over digit, so not one tab,
    dot or «20» of the template moves and every digit keeps its own box.

    Only a date that *is* the previous start or the previous end is written
    over. That is what keeps «Дата заключения договора 28.06.2024 г.» — the
    same date, a different field — out of it, and it is what makes each write
    all eight digits or none: a date the program cannot read whole cannot be
    matched either, so it is reported instead of half-changed.
    """
    lines = list(_strips(doc))
    previous = _previous_cover(lines)
    swap = dict(zip(previous, (start, end), strict=True)) if previous else {}
    for pieces in lines:
        text, where, fresh = _strip_text(pieces)
        for lo, hi, value in _cover_on(text, fresh):
            # one character for one, or the digits stop landing in their boxes
            if (new := swap.get(value)) is not None and len(new) == hi - lo:
                _put(pieces, where, lo, hi, new)
                report.filled["term"] = f"{start} — {end}"

    if not any(_names_the_term(pieces) for pieces in lines):
        return                      # this form does not ask for a срок at all
    left = any(value in swap for pieces in lines
               for _lo, _hi, value in _cover_dates(pieces))
    # a hole no boxes could fill: the day and the month are simply not on the
    # line, and the four digits that belong in them were not to be found
    hollow = any(_HOLLOW.search(_strip_text(pieces)[0]) and _spliced(pieces) is None
                 for pieces in lines)
    if previous is None or left or hollow:
        report.skipped.append(
            "Срок страхования — бланкада катак-катак ёзилган, қўлда тўлдиринг")


def _dates(doc, *, start: str, end: str, signed: str, report: Report) -> None:
    """Write the policy's own dates wherever the form prints them."""
    _cover(doc, start=start, end=end, report=report)

    for paragraph in _all_paragraphs(doc):
        runs = runs_of(paragraph)
        text = text_of(runs)
        if not text.strip() or not _SIGNED_ON.search(text):
            continue
        # only the date; «Дата заключения договора» and everything else on the
        # line belongs to the insurer's own wording
        edits = []
        if (m := _LONG_DATE.search(text)) is not None:
            edits.append((m.start(), m.end(), _long(signed)))
        if (m := _DMY.search(text)) is not None:
            edits.append((m.start(), m.end(), signed))
        for lo, hi, value in sorted(edits, reverse=True):
            replace_span(runs, lo, hi, value)
            report.filled["signed_on"] = signed


_MONTHS = ("января", "февраля", "марта", "апреля", "мая", "июня",
           "июля", "августа", "сентября", "октября", "ноября", "декабря")


def _long(dmy: str) -> str:
    """«28.07.2026» → «28» июля 2026 — the shape those templates print."""
    day, month, year = dmy.split(".")
    return f"«{day}» {_MONTHS[int(month) - 1]} {year}"


def fill(template: Path, out: Path, *, values: dict[str, str],
         drivers: list[tuple[str, str]], unlimited: bool,
         start: str = "", end: str = "", signed: str = "") -> Report:
    """Fill one insurer's policy template for one car and up to four drivers."""
    import docx

    doc = docx.Document(str(template))
    report = Report()
    _note_signature(doc, report)
    _paragraph_values(doc, values, report)
    _vin_grids(doc, values.get("vin", ""), report)
    # The previous car's own numbers are swapped first: wherever the VIN or the
    # plate already stands, that is where this one goes. Only if it is nowhere
    # does an empty labelled cell get it — otherwise it lands twice.
    _by_pattern(doc, values, report)
    placed = {key for key in ("vin", "plate", "sts") if key in report.filled}
    if "vin_grid" in report.filled:
        placed.add("vin")
    # A heading with the value in the row under it first, then a labelled cell
    # the template left blank — in that order, so a form that has both does not
    # end up with the value written twice.
    _column_values(doc, values, report, skip=frozenset(placed))
    _cell_gaps(doc, values, report, skip=frozenset(placed))
    if "sts" not in report.filled:
        _document_line(doc, values.get("sts", ""), report)
    _fill_drivers(doc, [] if unlimited else drivers, report)
    _choose(doc, unlimited, report)
    if start and end:
        _dates(doc, start=start, end=end, signed=signed or start, report=report)
    written = set(report.filled)
    if "vin_grid" in written:
        written.add("vin")     # it went into the boxes; telling the operator to
                               # write it by hand would have them write it twice
    for key in values:
        if key not in written and values[key]:
            report.skipped.append(f"{key} — бланкада бу жой топилмади, қўлда ёзинг")
    seen: list[str] = []
    for note in report.skipped:
        if note not in seen:
            seen.append(note)
    report.skipped = seen
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    log.info("ОСАГО filled from %s → %s (%d drivers)",
             template.name, out.name, report.drivers)
    return report
