"""Read the OLD worker's values out of a firm's own Word file.

Used when the office adds a NEW firm: the uploaded ТД/УВ is scanned the
same way the bundled ten were, so replacement works without any hand map.
"""

from __future__ import annotations

import re

import docx

DATE = r"(\d{2}\.\d{2}\.\d{4})"


def doc_texts(path):
    d = docx.Document(str(path))
    out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    out.extend(x.strip() for x in t.split("\n") if x.strip())
    return out


def uv_values(texts):
    """Label + value in ONE paragraph, or the value on the next line."""
    got, section = {}, ""
    plain = [("Фамилия (рус.)", "surname"), ("Имя (рус.)", "name"),
             ("Отчество (рус.)", "patronymic"),
             ("Дата заключения договора", "deal_date"),
             ("Дата рождения", "birth_date"), ("Пол", "gender"),
             ("Место рождения, населенный пункт", "birth_place"),
             ("Гражданство", "citizenship")]
    sectioned = [("Серия бланка", "pat_blank_series"),
                 ("Номер бланка", "pat_blank_number"),
                 ("Дата выдачи", "pass_issued"), ("Кем выдан",
                                                  "pass_issued_by"),
                 ("Серия", None), ("Номер", None)]
    i = 0
    while i < len(texts):
        t = texts[i]
        if t.startswith("Документ, удостоверяющий личность"):
            section = "pass"
        elif t.startswith("Сведения о разрешении на работу"):
            section = "pat"
        if t.startswith("Профессия, специальность"):
            rest = t.split("по договору", 1)
            value = rest[1].strip() if len(rest) > 1 and rest[1].strip() \
                else texts[i + 1]
            got["profession"] = value
            i += 1
            continue
        for label, key in plain + sectioned:
            if key is None:                    # Серия/Номер per section
                key = (f"{'pass' if section == 'pass' else 'pat'}_"
                       f"{'series' if label == 'Серия' else 'number'}")
            if section == "pat" and key in ("pass_issued", "pass_issued_by"):
                key = {"pass_issued": "pat_issued",
                       "pass_issued_by": "pat_issued_by"}[key]
            if key in got:
                continue
            if t == label and i + 1 < len(texts):
                got[key] = texts[i + 1]
                break
            if t.startswith(label + " "):
                value = t[len(label):].strip()
                if value:
                    got[key] = value
                    break
        i += 1
    got.pop("pat_issued", None)
    got.pop("pat_issued_by", None)
    return got


def td_values(texts):
    joined = "\n".join(texts)
    got = {}

    def take(key, pattern):
        if key not in got:
            m = re.search(pattern, joined)
            if m:
                got[key] = " ".join(m.group(1).split()).strip(" .,:;")

    take("surname", r"Фамилия\s*\(рус\.?\)\s*:?\s*([А-ЯЁ][а-яёА-ЯЁ]+)")
    take("name", r"Имя\s*\(рус\.?\)\s*:?\s*([А-ЯЁ][а-яёА-ЯЁ]+)")
    take("surname", r"Фамилия\t+\s*([А-ЯЁ][а-яёА-ЯЁ]+)")
    take("name", r"Имя\t+\s*([А-ЯЁ][а-яёА-ЯЁ]+)")
    take("patronymic", r"Отчество\t+\s*([А-ЯЁ][а-яёА-ЯЁ]+)")
    take("profession", r"на должность\s+([А-ЯЁ][а-яё]+(?:[ ]+[а-яё]+){0,3})")
    take("patronymic",
         r"Отчество\s*\(рус\.?\)\s*:?\s*([А-ЯЁ][а-яёА-ЯЁ]+(?:[ ][А-ЯЁа-яё]+)?)")
    take("fio", r"Работник\s*:\s*([А-ЯЁ][а-яёА-ЯЁ]+(?:[ ]+[А-ЯЁ][а-яёА-ЯЁ]+){1,3})")
    take("fio", r"стороны,\s*([А-ЯЁ][а-яёА-ЯЁ]+(?:[ ]+[А-ЯЁ][а-яёА-ЯЁ]+){1,3})\s*\(физическое")
    take("birth_date", r"Дата рождения\D{0,6}" + DATE)
    take("citizenship", r"Гражданство[.:\s]+([А-ЯЁ][а-яё]+)")
    take("pass_series", r"паспорт[^\n]*?Серия\s+([A-ZА-ЯЁ0-9]{1,4})\s")
    take("pass_number", r"паспорт[^\n]*?Номер\s*:?\s*(\S{6,})")
    take("pass_issued", r"Дата выдачи\s*:?\s+" + DATE)
    take("profession", r"должност[иь]\s*:\s*([А-ЯЁ][а-яё]+(?:[ ]+[а-яё]+){0,3})")
    take("pat_series", r"Патент[^\n]*?Серия\s+(\S{1,4})\s")
    take("pat_number", r"Патент[^\n]*?Номер\s+(\d{6,})")
    take("pat_series", r"патент\s*\(если требуется\s*\)\s*(\S+)\s+\d{6,}")
    take("pat_number", r"патент\s*\(если требуется\s*\)\s*\S+\s+(\d{6,})")
    take("pat_issued", r"патент[^\n]*?выдан\s+" + DATE)
    for t in texts[:14]:
        if re.fullmatch(DATE, t.strip()):
            got["deal_date"] = t.strip()
            break
    if "deal_date" not in got:
        m = re.search(r"^(?:г\.[^\n]{0,40})?\s*" + DATE + r"\s*$", joined, re.M)
        if m:
            got["deal_date"] = m.group(1)
    stops = {"Дата", "Гражданство", "Иностранный", "заключили", "настоящий"}
    for key in ("fio", "profession", "patronymic"):
        if key in got:
            words = got[key].split()
            while words and words[-1] in stops:
                words.pop()
            got[key] = " ".join(words)
    if "fio" not in got and {"surname", "name"} <= set(got):
        got["fio"] = " ".join([got["surname"], got["name"],
                               got.get("patronymic", "")]).strip()
    return got
