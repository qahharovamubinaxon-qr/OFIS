"""Word (.docx) трудовой/уведомление templates: swap the old worker for the new.

Pattern-based like the PDF trud editor — no fixed positions, so any firm's docx
works. The Госуслуги labels the worker's data always carries («Фамилия (рус.)»,
«Дата рождения», «Серия … Номер … Кем выдан») are handled by
:mod:`src.services.docx_worker`, which rewrites only the value beside each label
and so keeps its font. What that leaves — the «гражданин республики … именуемый»
clause, «Срок договора», the signature «ФАМИЛИЯ И.О.» — is picked up afterwards
by the older text heuristics, which never touch a paragraph the label pass
already handled.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from src.pdf.formatters import _date_dmy, _date_long_g
from src.services.docx_worker import (
    iter_paragraphs as _iter_paragraphs,
    swap_header_date,
    swap_worker,
)

_FIO_CLAUSE = re.compile(
    r"(гражданин(?:ка)?\s+республики\s+).+?(\s+именуем)", re.IGNORECASE | re.DOTALL
)
_DMY = re.compile(r"\d{2}\.\d{2}\.\d{4}")
_LONG_DATE = re.compile(r"^\d{1,2}\s+[а-яё]+(\s+\d{4}\s*г?\.?)?$", re.IGNORECASE)
_YEAR_ONLY = re.compile(r"^\d{4}\s*г\.?$")
_SIGN_ABBR = re.compile(r"^[А-ЯЁ]{2,}\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.$")


def _set_text(paragraph, text: str) -> None:
    """Replace a paragraph's text, keeping the first run's formatting.

    Removes EVERY inline child (runs inside hyperlinks/smartTags included —
    some templates wrap the worker's name in them, which ``paragraph.runs``
    does not expose), then rewrites via the first plain run."""
    first_run = paragraph.runs[0] if paragraph.runs else None
    for child in list(paragraph._p):
        if not child.tag.endswith("}pPr") and (first_run is None or child is not first_run._r):
            paragraph._p.remove(child)
    if first_run is not None:
        first_run.text = text
    else:
        paragraph.add_run(text)


class TrudDocxEditor:
    def fill_trudovoy(
        self,
        template: Path,
        out: Path,
        *,
        form_date: date,
        fio_upper: str,          # ШОДИЕВ БОКИ ЗОКИР УГЛИ
        citizenship_upper: str,  # УЗБЕКИСТАН
        sign_abbr: str,          # ШОДИЕВ Б.З.
        contract_end: date | None,
        profession: str,
        fio_title: str = "",         # Рахимов Бадриддин Нормуродович
        birth_date: date | None = None,
        passport_number: str = "",
        passport_issue: date | None = None,
        worker: dict[str, str] | None = None,
    ) -> Path:
        import docx

        doc = docx.Document(str(template))
        long_date = _date_long_g(form_date)

        # First the labelled values — any firm's contract writes them the same
        # way, and this keeps each value's own font.
        report = swap_worker(doc, worker or {})
        header = swap_header_date(doc, long_date.removesuffix(" г."))
        done = set(report.touched)
        if header is not None:
            done.add(header)

        srok_done = False
        bare_dates = 0  # 1st bare dd.mm.yyyy cell = start, 2nd = contract end
        for p in _iter_paragraphs(doc):
            if id(p._p) in done:
                continue
            text = p.text
            stripped = text.strip()
            if not stripped:
                continue
            if _FIO_CLAUSE.search(text):
                new = _FIO_CLAUSE.sub(
                    lambda m: f"{m.group(1)}{citizenship_upper} {fio_upper}{m.group(2)}", text
                )
                _set_text(p, new)
                continue
            if "Срок договора" in text and _DMY.search(text) and not srok_done:
                dates = [_date_dmy(form_date)]
                if contract_end:
                    dates.append(_date_dmy(contract_end))
                it = iter(dates)
                _set_text(p, _DMY.sub(lambda m: next(it, m.group(0)), text))
                srok_done = True
                continue
            if stripped.startswith("Работник:") and fio_title:
                _set_text(p, "Работник:\t" + fio_title)
                continue
            if stripped.startswith("Дата рожд") and birth_date:
                _set_text(p, _DMY.sub(_date_dmy(birth_date), text.rstrip()))
                continue
            if stripped.startswith("Паспорт") and passport_number:
                out_t = re.sub(r"\d{6,12}", passport_number, text.rstrip(), count=1)
                if passport_issue:
                    out_t = _DMY.sub(_date_dmy(passport_issue), out_t)
                _set_text(p, out_t)
                continue
            if "Срок договора" not in text and _DMY.fullmatch(stripped):
                bare_dates += 1
                if bare_dates == 2 and contract_end:
                    _set_text(p, _date_dmy(contract_end))
                else:
                    _set_text(p, _date_dmy(form_date))
                continue
            if _LONG_DATE.match(stripped) and "г" in stripped and len(stripped) > 8:
                _set_text(p, long_date)
                continue
            if _LONG_DATE.match(stripped) and len(stripped) <= 12 and any(
                mon in stripped.lower() for mon in
                ("январ", "феврал", "март", "апрел", "мая", "июн", "июл",
                 "август", "сентябр", "октябр", "ноябр", "декабр")
            ):
                _set_text(p, f"{form_date.day:02d} {_date_long_g(form_date).split()[1]}")
                continue
            if _YEAR_ONLY.match(stripped):
                _set_text(p, f"{form_date.year} г.")
                continue
            if _SIGN_ABBR.match(stripped):
                _set_text(p, sign_abbr)
                continue
            m = re.search(r"(должности:|обязанности:)\s*(\S.*)$", stripped)
            if m:
                _set_text(p, stripped[: stripped.index(m.group(1)) + len(m.group(1))] + " " + profession)
        # A separate «Срок договора» row may keep bare dd.mm.yyyy cells; handled
        # above via the fullmatch branch (first = start, others already equal).
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out

    def fill_uvedomlenie(self, template: Path, out: Path,
                         values: dict[str, str]) -> Path:
        """``values`` keys: surname, name, patronymic, birth_date, gender,
        citizenship, birth_place, pass_series, pass_number, pass_issue_date,
        pass_issued_by, pat_series, pat_number, region, blank_series,
        blank_number, profession, contract_date, work_address.

        The employer block above it (наименование, ОГРН, ИНН, КПП, ОКВЭД,
        телефон) is the firm's and is never touched.
        """
        import docx

        doc = docx.Document(str(template))
        swap_worker(doc, values)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out


_FIO_TITLE = re.compile(r"^[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(\s+[А-ЯЁ][а-яё]+)*(\s+[Уу]гли)?$")


class HodDocxEditor:
    """Ходатайство о переоформлении патента (firm-specific .docx form)."""

    def fill(self, template: Path, out: Path, *, fio_title: str,
             citizenship_upper: str, pat_series: str, pat_number: str,
             passport_line: str, form_date: date) -> Path:
        import docx

        doc = docx.Document(str(template))
        pat_number_done = False
        for p in _iter_paragraphs(doc):
            stripped = p.text.strip()
            if not stripped:
                continue
            if re.fullmatch(r"\d{2}", stripped) and pat_series:
                _set_text(p, pat_series)
            elif re.fullmatch(r"\d{9,11}", stripped) and not pat_number_done:
                _set_text(p, pat_number)
                pat_number_done = True  # the later ИНН cell stays untouched
            elif re.fullmatch(r"[А-ЯЁ]{4,}", stripped) and citizenship_upper:
                _set_text(p, citizenship_upper)
            elif "выдан" in stripped and re.search(r"\d{6,}", stripped):
                _set_text(p, passport_line)
            elif _FIO_TITLE.fullmatch(stripped) and len(stripped.split()) >= 3:
                _set_text(p, fio_title)
            elif _DMY.fullmatch(stripped):
                _set_text(p, _date_dmy(form_date))
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out


def docx_to_pdf(docx_path: Path) -> Path | None:
    """Convert via MS Word when available (Windows); else return None and the
    .docx itself is the deliverable."""
    pdf = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf))
        if pdf.exists():
            return pdf
    except Exception:  # noqa: BLE001 - no Word / not Windows
        pass
    # LibreOffice fallback (soffice on PATH)
    try:
        import subprocess

        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(docx_path.parent), str(docx_path)],
            capture_output=True, timeout=120, check=True,
        )
        if pdf.exists():
            return pdf
    except Exception:  # noqa: BLE001
        pass
    return None
