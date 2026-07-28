"""Building a firm's two Word templates from its requisites alone.

Not every firm hands over a Word file. When the office has only the requisites —
name, ИНН, КПП, ОГРН, address, district, МВД office, director, and the stamp as
a PNG — this writes the pair itself:

* ``uvedomlenie.docx`` — the notification of a concluded contract, carrying the
  firm's own block already filled and the worker's section left open;
* ``trudovoy.docx`` — a трудовой договор with the firm as employer.

Both are written with the same Госуслуги labels every other firm's documents
use, so :mod:`src.services.docx_worker` fills them per worker exactly as it
fills an uploaded template — no second code path, no per-firm study.
"""

from __future__ import annotations

from pathlib import Path

from src.common.errors import ValidationError
from src.common.logging import get_logger
from src.domain.firm_details import FirmDetails

log = get_logger(__name__)

UVED_NAME = "uvedomlenie.docx"
TRUD_NAME = "trudovoy.docx"

_FONT = "Times New Roman"
_BODY_PT = 11
_TAB_CM = 7.0
_STAMP_CM = 4.2

#: Every worker field the уведомление asks for, in the order the form prints
#: them. A label with no value is what :mod:`docx_worker` writes into.
_WORKER_SECTIONS: tuple[tuple[str, tuple[tuple[str, str], ...]], ...] = (
    ("Сведения об иностранном гражданине", (
        ("Фамилия (рус.)", ""), ("Имя (рус.)", ""), ("Отчество (рус.)", ""),
        ("Дата рождения", ""), ("Пол", ""), ("Гражданство", ""),
        ("Место рождения, населенный пункт", ""),
    )),
    ("Документ, удостоверяющий личность иностранного гражданина", (
        ("Вид документа", "Иностранный паспорт"),
        ("Серия", ""), ("Номер", ""), ("Дата выдачи", ""), ("Кем выдан", ""),
    )),
    ("Сведения о разрешении на работу или патенте", (
        ("Документ", "Патент"),
        ("Серия", ""), ("Номер", ""), ("Регион", ""),
        ("Серия бланка", ""), ("Номер бланка", ""),
    )),
    ("Сведения о трудовой деятельности", (
        ("Профессия, специальность, должность, вид трудовой деятельности "
         "по договору", ""),
        ("Вид договора", "Трудовой"),
        ("Дата заключения договора", ""),
        ("Адрес места работы", ""),
    )),
)


# ------------------------------------------------------------------ paper


def _document():
    import docx
    from docx.shared import Cm, Pt

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)
    style.paragraph_format.space_after = Pt(2)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(1.4)
    return doc


def _title(doc, text: str, *, size: int = 13):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def _heading(doc, text: str):
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    return p


def _field(doc, label: str, value: str = ""):
    """«Label⇥value» — the value in its own run, so a swap keeps its font.

    The run is added even when empty: that is the spot the worker's value goes
    into, and giving it its own run is what lets it come out unbolded and in
    the body font.
    """
    from docx.shared import Cm

    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(_TAB_CM))
    p.add_run(label)
    p.add_run("\t")
    p.add_run(value)
    return p


def _body(doc, text: str, *, first_line: float = 1.0):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(first_line)
    p.add_run(text)
    return p


def _firm_block(doc, firm: FirmDetails) -> None:
    _field(doc, "Полное наименование", firm.name)
    if firm.short_name:
        _field(doc, "Сокращённое наименование", firm.short_name)
    for label, value in firm.requisites():
        _field(doc, label, value)
    if firm.address:
        _field(doc, "Юридический адрес", firm.address)
    if firm.district:
        _field(doc, "Район (городской округ)", firm.district)
    if firm.phone:
        _field(doc, "Основной телефон", firm.phone)
    _field(doc, "Статус", firm.status_line)
    if firm.signatory:
        _field(doc, "Руководитель", f"{firm.signatory_position} {firm.signatory}")


# ------------------------------------------------------------ уведомление


def build_uvedomlenie(firm: FirmDetails, out: Path) -> Path:
    doc = _document()
    _title(doc, "УВЕДОМЛЕНИЕ")
    _title(doc, "о заключении трудового договора с иностранным гражданином",
           size=11)

    _heading(doc, "Сведения о работодателе")
    _firm_block(doc, firm)

    for heading, fields in _WORKER_SECTIONS:
        _heading(doc, heading)
        for label, value in fields:
            _field(doc, label, value)

    _heading(doc, "Выбор подразделения МВД России")
    _field(doc, "Наименование территориального органа МВД России "
                "на региональном уровне", firm.mvd_office)

    _signature(doc, firm)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# --------------------------------------------------------------- договор


def _worker_block(doc) -> None:
    """The same labels the уведомление uses, so one engine fills both."""
    _field(doc, "Фамилия (рус.)")
    _field(doc, "Имя (рус.)")
    _field(doc, "Отчество (рус.)")
    _field(doc, "Дата рождения")
    _field(doc, "Пол")
    _field(doc, "Гражданство")
    _field(doc, "Иностранный паспорт: Серия")
    _field(doc, "Номер")
    _field(doc, "Дата выдачи")
    _field(doc, "Кем выдан")
    _field(doc, "Патент: Серия")
    _field(doc, "Номер")
    _field(doc, "Профессия, специальность, должность:")
    _field(doc, "Адрес места жительства:")


_CLAUSES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("1. ПРЕДМЕТ ДОГОВОРА", (
        "1.1. Работодатель предоставляет Работнику работу по обусловленной"
        " настоящим договором трудовой функции, обеспечивает условия труда,"
        " предусмотренные трудовым законодательством Российской Федерации,"
        " своевременно и в полном размере выплачивает Работнику заработную"
        " плату, а Работник обязуется лично выполнять определённую настоящим"
        " договором трудовую функцию и соблюдать правила внутреннего трудового"
        " распорядка, действующие у Работодателя.",
        "1.2. Работа по настоящему договору является для Работника основной.",
        "1.3. Работник осуществляет трудовую деятельность на основании патента"
        " в пределах субъекта Российской Федерации, на территории которого"
        " патент выдан, и по указанной в патенте профессии.",
        "1.4. Настоящий договор вступает в силу со дня его подписания обеими"
        " Сторонами.",
    )),
    ("2. ПРАВА И ОБЯЗАННОСТИ РАБОТНИКА", (
        "2.1. Работник обязан добросовестно исполнять свои трудовые"
        " обязанности, соблюдать правила внутреннего трудового распорядка,"
        " трудовую дисциплину и требования по охране труда и обеспечению"
        " безопасности труда.",
        "2.2. Работник обязан бережно относиться к имуществу Работодателя и"
        " других работников, незамедлительно сообщать Работодателю о"
        " возникновении ситуации, представляющей угрозу жизни и здоровью людей"
        " или сохранности имущества.",
        "2.3. Работник обязан сообщать Работодателю об изменении сведений,"
        " указанных в разделе о Работнике настоящего договора, а также о"
        " продлении, переоформлении или прекращении действия патента.",
        "2.4. Работник имеет права, предусмотренные статьёй 21 Трудового"
        " кодекса Российской Федерации.",
    )),
    ("3. ПРАВА И ОБЯЗАННОСТИ РАБОТОДАТЕЛЯ", (
        "3.1. Работодатель обязан соблюдать трудовое законодательство"
        " Российской Федерации, предоставить Работнику работу, обусловленную"
        " настоящим договором, и обеспечить безопасные условия труда.",
        "3.2. Работодатель обязан выплачивать в полном размере причитающуюся"
        " Работнику заработную плату в установленные сроки.",
        "3.3. Работодатель обязан осуществлять обязательное социальное"
        " страхование Работника в порядке, установленном федеральными законами.",
        "3.4. Работодатель обязан уведомить территориальный орган федерального"
        " органа исполнительной власти в сфере внутренних дел о заключении и о"
        " прекращении настоящего договора в срок, не превышающий трёх рабочих"
        " дней с даты заключения или прекращения.",
        "3.5. Работодатель имеет права, предусмотренные статьёй 22 Трудового"
        " кодекса Российской Федерации.",
    )),
    ("4. ОПЛАТА ТРУДА", (
        "4.1. За выполнение трудовых обязанностей Работнику устанавливается"
        " заработная плата в размере ____________ рублей в месяц.",
        "4.2. Заработная плата выплачивается не реже чем каждые полмесяца, в"
        " дни, установленные правилами внутреннего трудового распорядка"
        " Работодателя.",
        "4.3. Из заработной платы Работника удерживаются налоги и иные"
        " обязательные платежи в порядке и размерах, предусмотренных"
        " законодательством Российской Федерации.",
    )),
    ("5. РАБОЧЕЕ ВРЕМЯ И ВРЕМЯ ОТДЫХА", (
        "5.1. Работнику устанавливается пятидневная рабочая неделя"
        " продолжительностью 40 часов с двумя выходными днями.",
        "5.2. Работнику предоставляется ежегодный оплачиваемый отпуск"
        " продолжительностью 28 календарных дней.",
        "5.3. Время начала и окончания работы, перерывы для отдыха и питания"
        " определяются правилами внутреннего трудового распорядка Работодателя.",
    )),
    ("6. ОТВЕТСТВЕННОСТЬ СТОРОН", (
        "6.1. Стороны несут ответственность за неисполнение или ненадлежащее"
        " исполнение обязательств по настоящему договору в соответствии с"
        " законодательством Российской Федерации.",
        "6.2. Работник несёт материальную ответственность за прямой"
        " действительный ущерб, причинённый Работодателю его виновными"
        " действиями (бездействием).",
    )),
    ("7. ИЗМЕНЕНИЕ И ПРЕКРАЩЕНИЕ ДОГОВОРА", (
        "7.1. Изменения и дополнения к настоящему договору оформляются"
        " дополнительным соглашением Сторон в письменной форме, являющимся"
        " неотъемлемой частью настоящего договора.",
        "7.2. Настоящий договор прекращается по основаниям, предусмотренным"
        " Трудовым кодексом Российской Федерации, в том числе в связи с"
        " окончанием срока действия патента Работника.",
    )),
    ("8. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", (
        "8.1. Во всём, что не предусмотрено настоящим договором, Стороны"
        " руководствуются законодательством Российской Федерации.",
        "8.2. Настоящий договор составлен в двух экземплярах, имеющих"
        " одинаковую юридическую силу, по одному для каждой из Сторон.",
        "8.3. Споры между Сторонами разрешаются путём переговоров, а при"
        " недостижении согласия — в порядке, установленном законодательством"
        " Российской Федерации.",
    )),
)


def build_trudovoy(firm: FirmDetails, out: Path, *, header_date: str) -> Path:
    """``header_date`` («01 января 2026») is a placeholder — every generated
    contract re-dates it to the day the worker's documents are made."""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Cm

    doc = _document()
    _title(doc, "ТРУДОВОЙ ДОГОВОР № ______")

    place = doc.add_paragraph()
    place.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
    place.add_run(firm.district or "г. Москва")
    place.add_run("\t")
    place.add_run(f"{header_date} года")

    _body(doc, f"{firm.acting_clause()}, именуемый в дальнейшем «Работодатель»,"
               " с одной стороны, и гражданин, сведения о котором указаны ниже,"
               " именуемый в дальнейшем «Работник», с другой стороны, заключили"
               " настоящий трудовой договор о нижеследующем:")

    _heading(doc, "СВЕДЕНИЯ О РАБОТНИКЕ")
    _worker_block(doc)

    for heading, clauses in _CLAUSES:
        _heading(doc, heading)
        for clause in clauses:
            _body(doc, clause)

    _heading(doc, "9. АДРЕСА, РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    _field(doc, "Работодатель", firm.name)
    for label, value in firm.requisites():
        _field(doc, label, value)
    if firm.address:
        _field(doc, "Юридический адрес", firm.address)
    if firm.phone:
        _field(doc, "Основной телефон", firm.phone)

    _signature(doc, firm, right="Работник")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _signature(doc, firm: FirmDetails, *, right: str = "") -> None:
    from docx.shared import Cm, Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(10.0))
    p.add_run(f"{firm.signatory_position} ______________ "
              f"{firm.signatory or '__________________'}")
    if right:
        p.add_run("\t")
        p.add_run(f"{right} ______________")

    if firm.stamp_path and firm.stamp_path.exists():
        stamp = doc.add_paragraph()
        stamp.paragraph_format.space_before = Pt(4)
        try:
            stamp.add_run().add_picture(str(firm.stamp_path), width=Cm(_STAMP_CM))
        except Exception as exc:  # noqa: BLE001 - a PNG Word cannot read
            log.warning("Stamp not embedded (%s): %s", firm.stamp_path, exc)
            doc.paragraphs[-1]._p.getparent().remove(doc.paragraphs[-1]._p)


# ------------------------------------------------------------------ both


def build(firm: FirmDetails, folder: Path, *, header_date: str) -> tuple[Path, Path]:
    """Write both templates into ``folder``. Returns (трудовой, уведомление)."""
    try:
        import docx  # noqa: F401
    except ImportError as exc:  # pragma: no cover - python-docx is a hard dep
        raise ValidationError("python-docx ўрнатилмаган") from exc

    trud = build_trudovoy(firm, folder / TRUD_NAME, header_date=header_date)
    uved = build_uvedomlenie(firm, folder / UVED_NAME)
    log.info("Built templates for %s in %s", firm.display_name, folder)
    return trud, uved
