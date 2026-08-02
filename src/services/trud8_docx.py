"""Swap one worker for another inside a firm's own Word file.

The firm's ТД/УВ ships as the .docx it was really typed in, together with
the OLD worker's values keyed in a json. Filling = replacing those strings
with the new worker's — Word keeps every font, underline and reflow itself,
which is exactly what printing onto a PDF could not give.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from src.common.errors import OfisError
from src.common.logging import get_logger

log = get_logger(__name__)


def _apply_to_runs(paragraph, old: str, new: str) -> int:
    """Replace across run boundaries, keeping each run's formatting."""
    made = 0
    search_from = 0
    while True:
        runs = paragraph.runs
        text = "".join(r.text for r in runs)
        idx = text.find(old, search_from)
        if idx < 0:
            return made
        pos = 0
        spans = []
        for r in runs:
            spans.append((pos, pos + len(r.text), r))
            pos += len(r.text)
        start, end = idx, idx + len(old)
        first = True
        for s, e, r in spans:
            if e <= start or s >= end:
                continue
            a, b = max(start, s) - s, min(end, e) - s
            r.text = r.text[:a] + (new if first else "") + r.text[b:]
            first = False
        search_from = idx + len(new)
        made += 1


def _paragraphs(doc):
    def walk(container):
        yield from container.paragraphs
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from walk(cell)

    yield from walk(doc)
    for section in doc.sections:
        yield from walk(section.header)
        yield from walk(section.footer)


def fill(template: Path, replacements: dict[str, str], target: Path) -> Path:
    """Copy ``template`` to ``target`` with every old value replaced."""
    import docx

    template = Path(template)
    if not template.exists():
        raise OfisError("Фирманинг Word бланкаси топилмади.")
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, target)
    doc = docx.Document(str(target))

    long_pairs = sorted(
        ((old, new) for old, new in replacements.items() if len(old) >= 4),
        key=lambda p: -len(p[0]))
    short_pairs = [(old, new) for old, new in replacements.items()
                   if 0 < len(old) < 4]
    total = 0
    for paragraph in _paragraphs(doc):
        for old, new in long_pairs:
            if old in paragraph.text:
                total += _apply_to_runs(paragraph, old, new)
        for old, new in short_pairs:
            # a two-letter серия must never touch dates or ИНН — it is only
            # swapped standing alone or right after its own label
            line = paragraph.text
            if line.strip() == old or re.search(
                    rf"(Серия( бланка)?|серия)\s+{re.escape(old)}(\s|$)",
                    line):
                total += _apply_to_runs(paragraph, old, new)
    doc.save(str(target))
    log.info("Word тўлдирилди: %s (%d алмаштириш)", target.name, total)
    return target


def to_pdf(source: Path) -> Path | None:
    """The same paper as PDF, through Word itself — None when Word is not
    there to ask (the бот then sends the .docx and says so)."""
    try:
        import comtypes.client
    except ImportError:
        try:
            from docx2pdf import convert
        except ImportError:
            return None
        out = source.with_suffix(".pdf")
        try:
            convert(str(source), str(out))
        except Exception:                      # noqa: BLE001
            return None
        return out if out.exists() else None
    out = source.with_suffix(".pdf")
    word = None
    try:
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        opened = word.Documents.Open(str(source))
        opened.SaveAs(str(out), FileFormat=17)
        opened.Close()
        return out if out.exists() else None
    except Exception:                          # noqa: BLE001
        return None
    finally:
        import contextlib

        if word is not None:
            with contextlib.suppress(Exception):
                word.Quit()
